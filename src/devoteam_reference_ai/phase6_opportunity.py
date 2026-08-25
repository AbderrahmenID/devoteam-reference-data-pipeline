from __future__ import annotations

import hashlib
import json
import os
import re
import unicodedata
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

import yaml


EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE)
PHONE_RE = re.compile(r"(?<!\w)(?:\+?\d[\d .()/-]{7,}\d)(?!\w)")
MONEY_RE = re.compile(
    r"(?<!\w)\d[\d\s.,]{1,18}\s*(?:EUR|USD|TND|DT|F\s*CFA|€|\$)(?!\w)",
    re.IGNORECASE,
)
ARABIC_RE = re.compile(r"[\u0600-\u06FF]")
SENTENCE_BOUNDARY_RE = re.compile(r"(?<=[.!?؟;؛])\s+")
BULLET_RE = re.compile(r"^\s*(?:[-*•▪◦‣]+|\d+[.)]|[A-Za-z][.)])\s*")
YEAR_RE = re.compile(r"\b(?:19|20)\d{2}\b")
SPACE_RE = re.compile(r"\s+")


class OpportunityAnalysisError(RuntimeError):
    """Raised when an opportunity cannot be processed safely."""


@dataclass(frozen=True)
class SourceSegment:
    text: str
    locator: str
    page_number_1_based: int | None


def utc_now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def normalize_text(value: Any) -> str:
    text = unicodedata.normalize("NFKD", str(value or ""))
    text = "".join(character for character in text if not unicodedata.combining(character))
    text = text.casefold().replace("’", "'")
    text = re.sub(r"[^\w\u0600-\u06ff]+", " ", text, flags=re.UNICODE)
    return SPACE_RE.sub(" ", text).strip()


def sanitize_opportunity_text(value: str) -> str:
    """Minimize personal and financial details before downstream analysis."""
    value = EMAIL_RE.sub("[EMAIL]", value)
    value = PHONE_RE.sub("[PHONE]", value)
    value = MONEY_RE.sub("[MONETARY_VALUE]", value)
    return SPACE_RE.sub(" ", value).strip()


def _atomic_text(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(value, encoding="utf-8")
    os.replace(temporary, path)


def _atomic_json(path: Path, value: Any) -> None:
    _atomic_text(path, json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True) + "\n")


def load_phase6_config(path: Path) -> dict:
    config = yaml.safe_load(path.read_text(encoding="utf-8"))
    if int(config.get("phase", 0)) != 6:
        raise ValueError("Expected Phase 6 configuration")
    security = config.get("security", {})
    if security.get("external_llm_enabled"):
        raise ValueError("Phase 6 defaults to deterministic local analysis; external LLM is blocked")
    if security.get("raw_opportunity_logging_allowed"):
        raise ValueError("Raw opportunity logging must remain disabled")
    if security.get("security_filters_disableable"):
        raise ValueError("Security filters may not be disableable")
    filters = config.get("filters", {})
    if not filters.get("human_confirmation_required"):
        raise ValueError("All proposed business filters require human confirmation")
    if filters.get("hidden_business_filters_allowed"):
        raise ValueError("Business filters must remain visible and editable")
    forbidden = {normalize_text(value) for value in filters.get("never_filter_fields", [])}
    if not {"project value", "team member", "personal data"}.issubset(forbidden):
        raise ValueError("The Phase 0 forbidden-filter baseline is incomplete")
    return config


def verify_phase5_dependency(phase5_root: Path, config: dict) -> dict:
    success_path = phase5_root / "_SUCCESS.json"
    manifest_path = phase5_root / "PHASE_5_MANIFEST.json"
    sums_path = phase5_root / "SHA256SUMS.txt"
    for path in (success_path, manifest_path, sums_path):
        if not path.exists():
            raise FileNotFoundError(f"Missing Phase 5 dependency: {path}")
    success = json.loads(success_path.read_text(encoding="utf-8"))
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    expected = config["input"]
    if sha256_file(manifest_path) != expected["expected_phase5_manifest_sha256"]:
        raise AssertionError("Phase 5 manifest hash changed")
    if sha256_file(sums_path) != expected["expected_phase5_sha256sums_sha256"]:
        raise AssertionError("Phase 5 SHA256SUMS hash changed")
    if success.get("manifest_sha256") != expected["expected_phase5_manifest_sha256"]:
        raise AssertionError("Phase 5 success marker does not match its manifest")
    if manifest.get("snapshot_id") != expected["snapshot_id"]:
        raise AssertionError("Phase 5 snapshot changed")
    if manifest.get("status") != "TECHNICAL_PASS" or manifest.get("qa_gate") != "PASS":
        raise AssertionError("Phase 5 technical gate is not complete")
    if int(manifest.get("external_llm_calls", -1)) != 0:
        raise AssertionError("Unexpected external LLM calls in Phase 5")
    if int(manifest.get("raw_user_query_log_rows", -1)) != 0:
        raise AssertionError("Unexpected raw-query logs in Phase 5")
    return {"success": success, "manifest": manifest}


def extract_opportunity_segments(path: Path, config: dict) -> list[SourceSegment]:
    path = path.resolve()
    if not path.is_file():
        raise FileNotFoundError(path)
    maximum = int(config["input"]["maximum_input_bytes"])
    if path.stat().st_size > maximum:
        raise OpportunityAnalysisError(f"Input exceeds {maximum} bytes")
    suffix = path.suffix.casefold()
    supported = {value.casefold() for value in config["input"]["supported_extensions"]}
    if suffix not in supported:
        raise OpportunityAnalysisError(f"Unsupported opportunity format: {suffix}")
    segments: list[SourceSegment] = []
    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8-sig", errors="strict")
        segments.append(SourceSegment(text=text, locator="text:1", page_number_1_based=None))
    elif suffix == ".pdf":
        import fitz

        with fitz.open(path) as document:
            for index, page in enumerate(document, start=1):
                segments.append(
                    SourceSegment(
                        text=page.get_text("text", sort=True),
                        locator=f"page:{index}",
                        page_number_1_based=index,
                    )
                )
        nonempty = sum(len(normalize_text(segment.text)) for segment in segments)
        if nonempty < int(config["input"]["minimum_extracted_characters"]):
            raise OpportunityAnalysisError(
                "PDF has insufficient digital text. Route this opportunity through the approved OCR path."
            )
    elif suffix == ".docx":
        from docx import Document

        document = Document(path)
        ordinal = 0
        for paragraph in document.paragraphs:
            ordinal += 1
            if paragraph.text.strip():
                segments.append(
                    SourceSegment(
                        text=paragraph.text,
                        locator=f"paragraph:{ordinal}",
                        page_number_1_based=None,
                    )
                )
        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if text:
                    segments.append(
                        SourceSegment(
                            text=text,
                            locator=f"table:{table_index}:row:{row_index}",
                            page_number_1_based=None,
                        )
                    )
    if not segments or not any(normalize_text(segment.text) for segment in segments):
        raise OpportunityAnalysisError("Opportunity input contains no usable text")
    return segments


def detect_language(text: str) -> str:
    normalized = normalize_text(text)
    if not normalized:
        return "unknown"
    arabic_count = len(ARABIC_RE.findall(text))
    latin_count = len(re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ]", text))
    if arabic_count > max(8, latin_count * 0.35):
        return "ar"
    words = set(normalized.split())
    french = {"le", "la", "les", "des", "pour", "avec", "doit", "devra", "mission", "référence", "reference"}
    english = {"the", "and", "for", "with", "must", "shall", "project", "reference", "experience"}
    fr_score = len(words & french)
    en_score = len(words & english)
    if fr_score == en_score == 0:
        return "mixed"
    return "fr" if fr_score >= en_score else "en"


def _units_from_segment(segment: SourceSegment, minimum_characters: int) -> Iterable[dict]:
    ordinal = 0
    for raw_line in segment.text.replace("\r", "\n").split("\n"):
        line = BULLET_RE.sub("", raw_line).strip()
        if not line:
            continue
        sentences = SENTENCE_BOUNDARY_RE.split(line) if len(line) > 280 else [line]
        for sentence in sentences:
            sentence = sanitize_opportunity_text(sentence)
            if len(normalize_text(sentence)) < minimum_characters:
                continue
            ordinal += 1
            yield {
                "text": sentence,
                "source_locator": f"{segment.locator}:unit:{ordinal}",
                "page_number_1_based": segment.page_number_1_based,
            }


def _pattern_matches(text: str, patterns: Iterable[str]) -> bool:
    normalized = f" {normalize_text(text)} "
    return any(f" {normalize_text(pattern)} " in normalized for pattern in patterns)


def classify_requirement(text: str, config: dict) -> str:
    rules = config["requirement_rules"]
    for label in ("MUST", "SHOULD", "PREFERRED"):
        if _pattern_matches(text, rules[label]):
            return label
    return "CONTEXT"


def extract_requirements(
    segments: list[SourceSegment], source_sha256: str, config: dict
) -> list[dict]:
    requirements: list[dict] = []
    seen: set[str] = set()
    minimum = int(config["analysis"]["minimum_unit_characters"])
    maximum = int(config["analysis"]["maximum_unit_characters"])
    for segment in segments:
        for unit in _units_from_segment(segment, minimum):
            text = unit["text"][:maximum].strip()
            normalized = normalize_text(text)
            if normalized in seen:
                continue
            seen.add(normalized)
            classification = classify_requirement(text, config)
            identity = f"{source_sha256}|{unit['source_locator']}|{normalized}"
            requirements.append(
                {
                    "requirement_id": f"REQ-{sha256_text(identity)[:12]}",
                    "requirement_text": text,
                    "language": detect_language(text),
                    "classification": classification,
                    "source_locator": unit["source_locator"],
                    "page_number_1_based": unit["page_number_1_based"],
                    "status": "PROPOSED",
                    "human_review_required": True,
                    "approved": False,
                }
            )
    if not requirements:
        raise OpportunityAnalysisError("No analyzable opportunity units were extracted")
    requirements.sort(key=lambda row: row["source_locator"])
    return requirements


def _value_matches(text_normalized: str, value: str, minimum_length: int) -> bool:
    normalized = normalize_text(value)
    if len(normalized) < minimum_length or normalized in {"oui", "non", "true", "false", "n a", "na"}:
        return False
    return f" {normalized} " in f" {text_normalized} "


def _canonical_alias_matches(text_normalized: str, values: Iterable[str], aliases: dict) -> list[str]:
    normalized_values = {normalize_text(value): value for value in values}
    output: list[str] = []
    for alias, canonical in aliases.items():
        if f" {normalize_text(alias)} " not in f" {text_normalized} ":
            continue
        actual = normalized_values.get(normalize_text(canonical))
        if actual:
            output.append(actual)
    return output


def compile_filter_proposals(
    requirements: list[dict], filter_values: dict, config: dict
) -> list[dict]:
    settings = config["filters"]
    supported = settings["supported_fields"]
    minimum = int(settings["minimum_value_characters"])
    hard_fields = set(settings["hard_candidate_fields"])
    proposals: list[dict] = []
    seen: set[tuple[str, str, str]] = set()
    negative_patterns = settings["exclusion_markers"]
    aliases = settings.get("value_aliases", {})
    for requirement in requirements:
        text_normalized = normalize_text(requirement["requirement_text"])
        exclusion = _pattern_matches(requirement["requirement_text"], negative_patterns)
        years = [int(value) for value in YEAR_RE.findall(requirement["requirement_text"])]
        experience_context = _pattern_matches(
            requirement["requirement_text"], settings["year_experience_markers"]
        )
        has_year_range_marker = _pattern_matches(
            requirement["requirement_text"],
            [*settings["year_after_markers"], *settings["year_before_markers"]],
        )
        for field in supported:
            # A phrase such as "depuis 2019" means a range, not the exact year
            # 2019. Emit only year_after/year_before below to avoid two
            # contradictory controls for the same evidence span.
            if field == "project_year" and years and experience_context and has_year_range_marker:
                continue
            values = list((filter_values.get(field) or {}).keys())
            matched = [value for value in values if _value_matches(text_normalized, value, minimum)]
            matched.extend(_canonical_alias_matches(text_normalized, values, aliases.get(field, {})))
            for value in sorted(set(matched), key=lambda item: normalize_text(item)):
                if exclusion:
                    behavior = "EXCLUSION_CANDIDATE"
                elif requirement["classification"] == "MUST" and field in hard_fields:
                    behavior = "HARD_CANDIDATE"
                elif requirement["classification"] in {"SHOULD", "PREFERRED"}:
                    behavior = "SOFT_PREFERENCE"
                else:
                    behavior = "CONTEXT_ONLY"
                key = (field, normalize_text(value), behavior)
                if key in seen:
                    continue
                seen.add(key)
                identity = f"{requirement['requirement_id']}|{field}|{normalize_text(value)}|{behavior}"
                proposals.append(
                    {
                        "filter_id": f"FLT-{sha256_text(identity)[:12]}",
                        "field": field,
                        "value": value,
                        "proposed_behavior": behavior,
                        "source_requirement_id": requirement["requirement_id"],
                        "evidence_text": requirement["requirement_text"],
                        "visible_to_user": True,
                        "requires_human_confirmation": True,
                        "confirmed": False,
                        "status": "PROPOSED",
                    }
                )
        for year in years if experience_context else []:
            if _pattern_matches(requirement["requirement_text"], settings["year_after_markers"]):
                field = "year_after"
            elif _pattern_matches(requirement["requirement_text"], settings["year_before_markers"]):
                field = "year_before"
            else:
                field = "project_year"
            behavior = (
                "HARD_CANDIDATE"
                if requirement["classification"] == "MUST" and field in hard_fields
                else "SOFT_PREFERENCE"
                if requirement["classification"] in {"SHOULD", "PREFERRED"}
                else "CONTEXT_ONLY"
            )
            key = (field, str(year), behavior)
            if key in seen:
                continue
            seen.add(key)
            identity = f"{requirement['requirement_id']}|{field}|{year}|{behavior}"
            proposals.append(
                {
                    "filter_id": f"FLT-{sha256_text(identity)[:12]}",
                    "field": field,
                    "value": year,
                    "proposed_behavior": behavior,
                    "source_requirement_id": requirement["requirement_id"],
                    "evidence_text": requirement["requirement_text"],
                    "visible_to_user": True,
                    "requires_human_confirmation": True,
                    "confirmed": False,
                    "status": "PROPOSED",
                }
            )
    proposals.sort(key=lambda row: (row["source_requirement_id"], row["field"], str(row["value"])))
    validate_filter_proposals(proposals, config)
    return proposals


def validate_filter_proposals(proposals: list[dict], config: dict) -> None:
    settings = config["filters"]
    supported = set(settings["supported_fields"]) | {"year_before", "year_after"}
    forbidden = {normalize_text(value) for value in settings["never_filter_fields"]}
    allowed_behaviors = {"HARD_CANDIDATE", "SOFT_PREFERENCE", "CONTEXT_ONLY", "EXCLUSION_CANDIDATE"}
    for proposal in proposals:
        if proposal["field"] not in supported:
            raise AssertionError(f"Unsupported proposed filter: {proposal['field']}")
        if normalize_text(proposal["field"]) in forbidden:
            raise AssertionError("Forbidden business filter proposed")
        if proposal["proposed_behavior"] not in allowed_behaviors:
            raise AssertionError("Unknown filter behavior")
        if not proposal["visible_to_user"] or not proposal["requires_human_confirmation"]:
            raise AssertionError("Every business filter must be visible and human-confirmed")
        if proposal["confirmed"] or proposal["status"] != "PROPOSED":
            raise AssertionError("Phase 6 may propose but may not silently confirm business filters")


def build_retrieval_query(requirements: list[dict], maximum_characters: int = 1800) -> str:
    ordered = sorted(
        requirements,
        key=lambda row: ({"MUST": 0, "SHOULD": 1, "PREFERRED": 2, "CONTEXT": 3}[row["classification"]], row["source_locator"]),
    )
    selected: list[str] = []
    length = 0
    for row in ordered:
        text = row["requirement_text"].strip()
        if length + len(text) + 1 > maximum_characters:
            continue
        selected.append(text)
        length += len(text) + 1
    return "\n".join(selected)


def analyze_opportunity(input_path: Path, filter_values_path: Path, config: dict) -> dict:
    expected_hash = config["input"]["expected_filter_values_sha256"]
    if sha256_file(filter_values_path) != expected_hash:
        raise AssertionError("Phase 4 filter-values hash changed")
    source_hash = sha256_file(input_path)
    segments = extract_opportunity_segments(input_path, config)
    requirements = extract_requirements(segments, source_hash, config)
    filter_values = json.loads(filter_values_path.read_text(encoding="utf-8"))
    proposals = compile_filter_proposals(requirements, filter_values, config)
    counts = {label: sum(row["classification"] == label for row in requirements) for label in ("MUST", "SHOULD", "PREFERRED", "CONTEXT")}
    return {
        "schema_version": 1,
        "phase": 6,
        "pipeline_version": config["pipeline_version"],
        "analysis_mode": config["analysis"]["mode"],
        "opportunity_id": f"OPP-{source_hash[:16]}",
        "source_file_name": input_path.name,
        "source_sha256": source_hash,
        "source_size_bytes": input_path.stat().st_size,
        "source_was_copied": False,
        "source_language": detect_language(" ".join(segment.text for segment in segments)),
        "external_llm_calls": 0,
        "raw_opportunity_log_rows": 0,
        "requirements": requirements,
        "requirement_counts": counts,
        "filter_proposals": proposals,
        "retrieval_query": build_retrieval_query(requirements),
        "business_filters_auto_applied": 0,
        "human_review_required": True,
        "status": "READY_FOR_HUMAN_REVIEW",
    }


def create_review_workbook(path: Path, analysis: dict) -> None:
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.worksheet.datavalidation import DataValidation

    workbook = openpyxl.Workbook()
    overview = workbook.active
    overview.title = "Overview"
    overview.append(["Control", "Value"])
    for row in (
        ("Opportunity ID", analysis["opportunity_id"]),
        ("Source file", analysis["source_file_name"]),
        ("Source SHA-256", analysis["source_sha256"]),
        ("Analysis mode", analysis["analysis_mode"]),
        ("External LLM calls", analysis["external_llm_calls"]),
        ("Requirements", len(analysis["requirements"])),
        ("Filter proposals", len(analysis["filter_proposals"])),
        ("Status", analysis["status"]),
        ("Rule", "Review every proposed requirement and filter; nothing is auto-approved."),
    ):
        overview.append(row)
    requirements = workbook.create_sheet("Requirements")
    requirements.append(
        [
            "requirement_id",
            "classification",
            "language",
            "requirement_text",
            "source_locator",
            "page_number_1_based",
            "reviewer_decision",
            "reviewer_correction",
            "reviewer_notes",
        ]
    )
    for item in analysis["requirements"]:
        requirements.append(
            [
                item["requirement_id"],
                item["classification"],
                item["language"],
                item["requirement_text"],
                item["source_locator"],
                item["page_number_1_based"],
                "PENDING",
                "",
                "",
            ]
        )
    filters = workbook.create_sheet("Filters")
    filters.append(
        [
            "filter_id",
            "field",
            "value",
            "proposed_behavior",
            "source_requirement_id",
            "evidence_text",
            "reviewer_decision",
            "reviewer_value",
            "reviewer_behavior",
            "reviewer_notes",
        ]
    )
    for item in analysis["filter_proposals"]:
        filters.append(
            [
                item["filter_id"],
                item["field"],
                item["value"],
                item["proposed_behavior"],
                item["source_requirement_id"],
                item["evidence_text"],
                "PENDING",
                "",
                "",
                "",
            ]
        )
    decision = DataValidation(type="list", formula1='"PENDING,APPROVE,REJECT,EDIT"')
    behavior = DataValidation(
        type="list",
        formula1='"HARD_CANDIDATE,SOFT_PREFERENCE,CONTEXT_ONLY,EXCLUSION_CANDIDATE"',
    )
    requirements.add_data_validation(decision)
    decision.add(f"G2:G{max(2, requirements.max_row)}")
    filter_decision = DataValidation(type="list", formula1='"PENDING,APPROVE,REJECT,EDIT"')
    filters.add_data_validation(filter_decision)
    filter_decision.add(f"G2:G{max(2, filters.max_row)}")
    filters.add_data_validation(behavior)
    behavior.add(f"I2:I{max(2, filters.max_row)}")
    for sheet in workbook.worksheets:
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="5B2C83")
            cell.alignment = Alignment(wrap_text=True)
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
    overview.column_dimensions["A"].width = 28
    overview.column_dimensions["B"].width = 95
    for index, width in enumerate([22, 17, 12, 80, 30, 22, 20, 55, 45], start=1):
        requirements.column_dimensions[openpyxl.utils.get_column_letter(index)].width = width
    for index, width in enumerate([22, 22, 28, 24, 24, 75, 20, 28, 24, 45], start=1):
        filters.column_dimensions[openpyxl.utils.get_column_letter(index)].width = width
    path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(path)


def _report(analysis: dict, input_mode: str) -> str:
    counts = analysis["requirement_counts"]
    return f"""# Phase 6 — Opportunity analysis report

**Status:** {analysis['status']}
**Input mode:** {input_mode}
**Opportunity ID:** `{analysis['opportunity_id']}`
**Analysis mode:** `{analysis['analysis_mode']}`

## Result

- Requirements: {len(analysis['requirements'])}
- MUST: {counts['MUST']}
- SHOULD: {counts['SHOULD']}
- PREFERRED: {counts['PREFERRED']}
- CONTEXT: {counts['CONTEXT']}
- Visible filter proposals: {len(analysis['filter_proposals'])}
- Business filters automatically applied: 0
- External LLM calls: 0
- Original opportunity copied: NO

## Boundary

This technical run proposes requirements and filters. It does not silently approve
or enforce them. A business user must review the workbook before the analysis can
drive a production search. The deterministic analyzer is a secure baseline, not a
claim that LLM-quality extraction has been validated.
"""


def write_analysis_run(
    *,
    input_path: Path,
    filter_values_path: Path,
    output_root: Path,
    config: dict,
    input_mode: str,
) -> tuple[Path, dict]:
    analysis = analyze_opportunity(input_path, filter_values_path, config)
    run_root = output_root / analysis["opportunity_id"] / config["pipeline_version"]
    success_path = run_root / config["output"]["success_marker"]
    if success_path.exists():
        verify_phase6_run(run_root, config)
        manifest = json.loads((run_root / "PHASE_6_MANIFEST.json").read_text(encoding="utf-8"))
        return run_root, manifest
    run_root.mkdir(parents=True, exist_ok=True)
    requirements_path = run_root / "requirements.jsonl"
    _atomic_text(
        requirements_path,
        "".join(json.dumps(row, ensure_ascii=False, sort_keys=True) + "\n" for row in analysis["requirements"]),
    )
    _atomic_json(run_root / "filter_proposals.json", analysis["filter_proposals"])
    _atomic_json(
        run_root / "opportunity_analysis.json",
        {key: value for key, value in analysis.items() if key not in {"requirements", "filter_proposals"}},
    )
    create_review_workbook(run_root / "OPPORTUNITY_REVIEW.xlsx", analysis)
    _atomic_text(run_root / "PHASE_6_REPORT.md", _report(analysis, input_mode))
    manifest = {
        "schema_version": 1,
        "phase": 6,
        "pipeline_version": config["pipeline_version"],
        "completed_at_utc": utc_now(),
        "status": "TECHNICAL_PASS_READY_FOR_HUMAN_REVIEW",
        "input_mode": input_mode,
        "opportunity_id": analysis["opportunity_id"],
        "source_file_name": analysis["source_file_name"],
        "source_sha256": analysis["source_sha256"],
        "source_was_copied": False,
        "phase5_manifest_sha256": config["input"]["expected_phase5_manifest_sha256"],
        "filter_values_sha256": config["input"]["expected_filter_values_sha256"],
        "requirements": len(analysis["requirements"]),
        "filter_proposals": len(analysis["filter_proposals"]),
        "business_filters_auto_applied": 0,
        "external_llm_calls": 0,
        "raw_opportunity_log_rows": 0,
        "human_review_required": True,
        "production_promotion_allowed": False,
    }
    manifest_path = run_root / "PHASE_6_MANIFEST.json"
    _atomic_json(manifest_path, manifest)
    hashed = [
        requirements_path,
        run_root / "filter_proposals.json",
        run_root / "opportunity_analysis.json",
        run_root / "OPPORTUNITY_REVIEW.xlsx",
        run_root / "PHASE_6_REPORT.md",
        manifest_path,
    ]
    sums = "".join(
        f"{sha256_file(path)}  {path.name}\n" for path in sorted(hashed, key=lambda item: item.name)
    )
    sums_path = run_root / "SHA256SUMS.txt"
    _atomic_text(sums_path, sums)
    success = {
        "status": "COMPLETE_REPRODUCIBLE_OPPORTUNITY_ANALYSIS",
        "created_at_utc": utc_now(),
        "pipeline_version": config["pipeline_version"],
        "opportunity_id": analysis["opportunity_id"],
        "manifest_sha256": sha256_file(manifest_path),
        "sha256sums_sha256": sha256_file(sums_path),
    }
    _atomic_json(success_path, success)
    verify_phase6_run(run_root, config)
    return run_root, manifest


def verify_phase6_run(run_root: Path, config: dict) -> dict:
    success_path = run_root / config["output"]["success_marker"]
    manifest_path = run_root / "PHASE_6_MANIFEST.json"
    sums_path = run_root / "SHA256SUMS.txt"
    for path in (success_path, manifest_path, sums_path):
        if not path.exists():
            raise FileNotFoundError(path)
    success = json.loads(success_path.read_text(encoding="utf-8"))
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    if success["manifest_sha256"] != sha256_file(manifest_path):
        raise AssertionError("Phase 6 manifest hash mismatch")
    if success["sha256sums_sha256"] != sha256_file(sums_path):
        raise AssertionError("Phase 6 checksum-file hash mismatch")
    for line in sums_path.read_text(encoding="utf-8").splitlines():
        expected, name = line.split("  ", 1)
        if sha256_file(run_root / name) != expected:
            raise AssertionError(f"Phase 6 output changed: {name}")
    if manifest.get("external_llm_calls") != 0 or manifest.get("business_filters_auto_applied") != 0:
        raise AssertionError("Phase 6 security or human-control boundary failed")
    if not manifest.get("human_review_required") or manifest.get("production_promotion_allowed"):
        raise AssertionError("Phase 6 promotion boundary failed")
    return {"success": success, "manifest": manifest}


def create_redacted_sample(path: Path) -> None:
    sample = """Mission de conseil pour la définition et la mise en œuvre d'une stratégie digitale bancaire en Tunisie.
Le prestataire doit démontrer au moins une référence dans le secteur Banque depuis 2019.
Il devra proposer une démarche de gouvernance, une architecture cible et une feuille de route de transformation.
Une expérience avec une institution financière en Afrique est souhaitée.
La capacité à produire des livrables en français et en anglais constitue un atout.
Les références sans document justificatif doivent être exclues de la sélection finale.
"""
    if not path.exists():
        _atomic_text(path, sample)
