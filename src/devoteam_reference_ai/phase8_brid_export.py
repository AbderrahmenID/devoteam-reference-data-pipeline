from __future__ import annotations

import hashlib
import json
import math
import os
import re
import shutil
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import yaml


STATUS = "TECHNICAL_PASS_AUDITED_EXPORT_COMPLETE"
LABEL = "SYNTHETIC_TEST_ONLY"

ROLE_LABELS_FR = {
    "SDSI_BFSI": "SDSI bancaire",
    "SDSI": "SDSI",
    "PCA": "Continuité d'activité / PCA",
    "SECURITY_PSSI": "Cybersécurité / PSSI",
    "SECURITY_ISO27001": "SMSI / ISO 27001",
    "SDSI_IMPLEMENTATION_AMOA": "AMOA / PMO de mise en œuvre SDSI",
    "DATA_GOVERNANCE_CLIENT_MASTER": "Gouvernance des données",
    "API_MANAGEMENT": "API management",
}

SECTOR_LABELS_FR = {
    "BANKING": "Banque",
    "INDUSTRY": "Industrie",
    "INSURANCE": "Assurance",
    "PUBLIC_ENTERPRISE": "Entreprise publique",
    "MICROFINANCE": "Microfinance",
}

EVIDENCE_LABELS_FR = {
    "ATTESTATION_ORIGINAL": "Attestation originale",
    "ATTESTATION_SCAN": "Attestation scannée",
}

ELIGIBILITY_BASIS_FR = {
    "ELIG-SIGNED": "Les dix sources retenues sont des attestations de référence client signées.",
    "ELIG-SDSI-3": "Les preuves d'achèvement des trois références SDSI se situent entre 2019 et 2026.",
    "ELIG-SDSI-BFSI-2": "Deux des trois références SDSI qualifiantes concernent le secteur bancaire.",
    "ELIG-PCA-2": "Deux attestations prouvent directement des missions de continuité/PCA.",
    "ELIG-SEC-2": "Une référence PSSI et une référence SMSI/ISO 27001 satisfont la règle.",
    "ELIG-AMOA-1": "L'attestation PMO signée couvre la mise en œuvre du portefeuille de projets issu du SDSI.",
    "ELIG-AFRICA-2": "Les dix références sont situées en Afrique du Nord ou en Afrique de l'Ouest.",
}

MISSION_COPY = {
    "33722b351996597b7286de77c5885b6d6009d8a8928d6715db518a2f12ae9016": {
        "title": "Accompagnement à l'urbanisation et au schéma directeur du SI",
        "summary": (
            "Mission SDSI bancaire couvrant le cadrage, l'urbanisation, le scénario "
            "cible et le plan de mise en œuvre."
        ),
    },
    "61c34f37d9342605280f6db71565ef0e268150f9d64f3ff3730dfefcf9fda53a": {
        "title": "Élaboration de la feuille de route et de l'architecture SI",
        "summary": (
            "Mission bancaire d'architecture et de feuille de route SI. La période "
            "d'achèvement a été vérifiée en 2019."
        ),
    },
    "f5d4cea0983c8859b7ddc576b5bb40d7e7dee47575821bb756d4ab6e60d20d3b": {
        "title": "Élaboration du schéma directeur des systèmes d'information",
        "summary": (
            "Diagnostic du SI, définition de la cible et de l'organisation SI, puis "
            "construction du portefeuille de projets et de la feuille de route."
        ),
    },
    "95abaf3e9e68fe399991e6feb166817adc7a1bd96424a0390809865cff492b25": {
        "title": "Mise en œuvre d'un plan de continuité d'activité",
        "summary": (
            "Mise en œuvre d'un PCA avec solutions de repli et formalisation des "
            "procédures de continuité."
        ),
    },
    "09f4d3df1b5d65aa8a6a049d51f4e0b4a6496bfcad459fa94d8ab5b76bc8b15b": {
        "title": "Développement et mise en place du processus de continuité SI",
        "summary": (
            "Mission de continuité alignée sur ISO 22301. Aucun crédit de "
            "cybersécurité n'est revendiqué pour cette référence."
        ),
    },
    "8a7c1b4fdb5247728408e984c28e4fb74c828017645be2cca795ab6d6a522633": {
        "title": "Mission RSSI en régie - PSSI, risques et incidents",
        "summary": (
            "Mission RSSI couvrant explicitement la PSSI, l'analyse des risques, "
            "la gestion des incidents et le traitement des vulnérabilités."
        ),
    },
    "3a73a8da6d1e981a72d15b57047983e28f02a7572966ef6c318b23e793c146b0": {
        "title": "Mise en place d'un SMSI et préparation ISO/IEC 27001",
        "summary": (
            "Mise en place d'un système de management de la sécurité de "
            "l'information et préparation à la certification ISO/IEC 27001:2013."
        ),
    },
    "de852d9e359a33166dfae8c5699d5b6b012e475b29e0dafd82516bf1a440e22a": {
        "title": "Accompagnement PMO du portefeuille de projets issu du SDSI",
        "summary": (
            "Accompagnement PMO du portefeuille de projets produit par le SDSI, "
            "avec consolidation des statuts, suivi budgétaire et reporting."
        ),
    },
    "3bb9cd8288d417032e1925281f4aabc3c0ab283042c069c2869831a53b6a4cbb": {
        "title": "Fiabilisation des données et référentiel client unique",
        "summary": (
            "Fiabilisation des référentiels et mise en place d'un référentiel "
            "client unique avec dispositif de pérennisation de la qualité."
        ),
    },
    "e5d39000071b4c5c2f1170fcbce4de576b4fe8643cd125a1a82df9d0da3e82bd": {
        "title": "Mise en place d'une API Gateway et transfert de compétences",
        "summary": (
            "Mise en place d'une API Gateway, publication d'API et transfert de "
            "compétences, d'après l'attestation source auditée."
        ),
    },
}


class Phase8ExportError(RuntimeError):
    """Raised when the audited export contract is incomplete or inconsistent."""


def utc_now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and math.isnan(value):
        return ""
    return " ".join(str(value).replace("\u00a0", " ").split())


def _as_int(value: Any) -> int:
    if value in (None, ""):
        return 0
    return int(float(value))


def _atomic_text(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(value, encoding="utf-8")
    os.replace(temporary, path)


def _atomic_json(path: Path, value: Any) -> None:
    _atomic_text(
        path,
        json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
    )


def load_phase8_config(path: Path) -> dict:
    config = yaml.safe_load(path.read_text(encoding="utf-8"))
    if int(config.get("phase", 0)) != 8:
        raise ValueError("Expected Phase 8 configuration")
    if config.get("label") != LABEL:
        raise ValueError("The BRID release must remain SYNTHETIC_TEST_ONLY")
    generation = config.get("generation", {})
    blocked_generation = (
        "external_llm_enabled",
        "external_translation_enabled",
        "automatic_client_delivery",
        "automatic_tender_scoring",
        "source_documents_copied",
        "source_page_images_embedded",
        "production_promotion_allowed",
    )
    if any(bool(generation.get(field)) for field in blocked_generation):
        raise ValueError("Phase 8 generation boundary was weakened")
    if generation.get("mode") != "deterministic_audited_export":
        raise ValueError("Unexpected Phase 8 generation mode")
    security = config.get("security", {})
    if any(
        bool(security.get(field))
        for field in (
            "raw_opportunity_logging_allowed",
            "source_mutation_allowed",
            "hidden_content_generation_allowed",
        )
    ):
        raise ValueError("Phase 8 security boundary was weakened")
    return config


def _verify_file(
    path: Path,
    *,
    expected_sha256: str,
    expected_size: int,
    label: str,
) -> dict:
    if not path.exists():
        raise FileNotFoundError(path)
    size = path.stat().st_size
    digest = sha256_file(path)
    if size != int(expected_size):
        raise AssertionError(f"{label} size changed")
    if digest != expected_sha256:
        raise AssertionError(f"{label} hash changed")
    return {"path": str(path), "size_bytes": size, "sha256": digest}


def verify_reference_template(path: Path, config: dict) -> dict:
    from docx import Document

    inputs = config["inputs"]
    result = _verify_file(
        path,
        expected_sha256=inputs["reference_template_sha256"],
        expected_size=inputs["reference_template_size_bytes"],
        label="Reference template",
    )
    document = Document(path)
    if len(document.tables) < 18:
        raise AssertionError("Reference template table structure changed")
    summary = document.tables[0]
    summary_text = " ".join(
        cell.text for row in summary.rows[:2] for cell in row.cells
    )
    for label in (
        "Intitulé du projet",
        "Client",
        "Pays",
        "Période",
        "Thématiques clés",
    ):
        if label not in summary_text:
            raise AssertionError(f"Reference template label missing: {label}")
    detail_labels = {
        clean_text(row.cells[0].text) for row in document.tables[1].rows
    }
    for label in (
        "Nom de la mission",
        "Pays",
        "Nom de l’Autorité Contractante",
        "Date de démarrage",
    ):
        if label not in detail_labels:
            raise AssertionError(f"Reference detail label missing: {label}")
    return {
        **result,
        "tables": len(document.tables),
        "sections": len(document.sections),
        "reference_slots": len(document.tables) - 1,
        "structure_verified": True,
    }


def verify_audit_workbook(path: Path, config: dict) -> dict:
    inputs = config["inputs"]
    return _verify_file(
        path,
        expected_sha256=inputs["audit_workbook_sha256"],
        expected_size=inputs["audit_workbook_size_bytes"],
        label="Audited Phase 7 workbook",
    )


def load_approval(path: Path, config: dict, audit_sha256: str) -> dict:
    approval = json.loads(path.read_text(encoding="utf-8"))
    expected = config["approval"]
    if approval.get("approval_status") != expected["required_status"]:
        raise Phase8ExportError("The audited shortlist is not approved")
    if approval.get("approval_scope") != expected["required_scope"]:
        raise Phase8ExportError("The approval scope does not match Phase 8")
    if approval.get("approval_statement") != expected["expected_statement"]:
        raise Phase8ExportError("The approval statement changed")
    if approval.get("source_workbook_sha256") != audit_sha256:
        raise Phase8ExportError("Approval is not bound to the audited workbook")
    if approval.get("label") != LABEL:
        raise Phase8ExportError("Approval label changed")
    if approval.get("approver_identity_recorded") is not False:
        raise Phase8ExportError("No approver identity may be invented")
    return approval


def _sheet_rows(sheet) -> list[dict]:
    values = list(sheet.iter_rows(values_only=True))
    if not values:
        return []
    headers = [clean_text(value) for value in values[0]]
    return [
        dict(zip(headers, row))
        for row in values[1:]
        if any(value not in (None, "") for value in row)
    ]


def _requirement_rows_for_reference(
    shortlist_row: dict,
    coverage_rows: list[dict],
) -> list[dict]:
    client = clean_text(shortlist_row["client"])
    role = clean_text(shortlist_row["verified_primary_role"])
    selected = []
    for row in coverage_rows:
        if clean_text(row.get("audit_status")) != "COVERED":
            continue
        code = clean_text(row.get("code"))
        clients = {
            item.strip()
            for item in clean_text(row.get("supporting_clients")).split("|")
            if item.strip()
        }
        if client not in clients:
            continue
        if client == "SUNU Assurance":
            if code == "SCOPE-DATA" and role != "DATA_GOVERNANCE_CLIENT_MASTER":
                continue
            if code == "SCOPE-API" and role != "API_MANAGEMENT":
                continue
            if code in {"SCOPE-DATA", "SCOPE-API"}:
                pass
            elif code not in {"CTX-SECTOR", "CTX-REGION"}:
                continue
        selected.append(
            {
                "code": code,
                "classification": clean_text(row.get("classification")),
                "requirement_text": clean_text(row.get("requirement_text")),
                "audit_status": "COVERED",
            }
        )
    return sorted(
        selected,
        key=lambda item: (
            0 if item["classification"] == "MUST" else 1,
            item["code"],
        ),
    )


def load_audited_case(path: Path, config: dict, approval: dict) -> dict:
    import openpyxl

    workbook = openpyxl.load_workbook(path, read_only=True, data_only=False)
    required_sheets = {
        "Audit Summary",
        "Audited Shortlist",
        "Audited Eligibility",
        "Audited Coverage",
        "Citation Adjudication",
        "Scoring Criteria",
    }
    missing = required_sheets - set(workbook.sheetnames)
    if missing:
        raise Phase8ExportError(f"Audit workbook sheets missing: {sorted(missing)}")

    summary_rows = _sheet_rows(workbook["Audit Summary"])
    summary = {
        clean_text(row.get("Control")): row.get("Value") for row in summary_rows
    }
    inputs = config["inputs"]
    if clean_text(summary.get("Label")) != LABEL:
        raise Phase8ExportError("Audit label changed")
    if (
        clean_text(summary.get("Source Phase 7 status"))
        != inputs["source_phase7_status"]
    ):
        raise Phase8ExportError("Unexpected source Phase 7 status")
    if clean_text(summary.get("Audit status")) != inputs["audit_status"]:
        raise Phase8ExportError("Unexpected evidence audit status")

    shortlist = _sheet_rows(workbook["Audited Shortlist"])
    expected = config["contracts"]
    if len(shortlist) != int(expected["selected_references"]):
        raise Phase8ExportError("Audited shortlist count changed")
    reference_ids = [clean_text(row.get("reference_id")) for row in shortlist]
    if len(reference_ids) != len(set(reference_ids)):
        raise Phase8ExportError("Audited shortlist contains duplicate reference IDs")
    if set(reference_ids) != set(MISSION_COPY):
        raise Phase8ExportError("Audited shortlist membership changed")
    if any(
        clean_text(row.get("citation_support")) != "SUPPORTED"
        for row in shortlist
    ):
        raise Phase8ExportError(
            "Every exported reference must have source-supported audit evidence"
        )
    if any(
        not clean_text(row.get("source_url")).startswith(
            "https://drive.google.com/file/d/"
        )
        or "#page=" not in clean_text(row.get("source_url"))
        for row in shortlist
    ):
        raise Phase8ExportError("Every exported reference needs an exact page URL")
    if any(
        clean_text(row.get("verified_primary_role")) not in ROLE_LABELS_FR
        for row in shortlist
    ):
        raise Phase8ExportError("Unknown audited reference role")

    eligibility = _sheet_rows(workbook["Audited Eligibility"])
    if len(eligibility) != int(expected["eligibility_rules_total"]):
        raise Phase8ExportError("Eligibility rule count changed")
    passed = [
        row
        for row in eligibility
        if clean_text(row.get("audit_status")) == "PASS"
        and _as_int(row.get("observed_count")) >= _as_int(row.get("required_count"))
    ]
    if len(passed) != int(expected["eligibility_rules_passed"]):
        raise Phase8ExportError("Audited eligibility gate failed")

    coverage = _sheet_rows(workbook["Audited Coverage"])
    must = [row for row in coverage if clean_text(row.get("classification")) == "MUST"]
    should = [
        row for row in coverage if clean_text(row.get("classification")) == "SHOULD"
    ]
    must_covered = [
        row for row in must if clean_text(row.get("audit_status")) == "COVERED"
    ]
    should_covered = [
        row for row in should if clean_text(row.get("audit_status")) == "COVERED"
    ]
    gaps = sorted(
        clean_text(row.get("code"))
        for row in coverage
        if clean_text(row.get("audit_status")) == "GAP"
    )
    checks = (
        (len(must), int(expected["must_requirements_total"]), "MUST total"),
        (
            len(must_covered),
            int(expected["must_requirements_covered"]),
            "MUST covered",
        ),
        (len(should), int(expected["should_requirements_total"]), "SHOULD total"),
        (
            len(should_covered),
            int(expected["should_requirements_covered"]),
            "SHOULD covered",
        ),
    )
    for actual, target, label in checks:
        if actual != target:
            raise Phase8ExportError(f"{label} changed: {actual} != {target}")
    if gaps != sorted(expected["explicit_gaps"]):
        raise Phase8ExportError("The explicit coverage gaps changed")

    adjudication = _sheet_rows(workbook["Citation Adjudication"])
    citation_stats = {
        "total": len(adjudication),
        "SUPPORTED": sum(
            clean_text(row.get("audit_status")) == "SUPPORTED"
            for row in adjudication
        ),
        "PARTIAL": sum(
            clean_text(row.get("audit_status")) == "PARTIAL"
            for row in adjudication
        ),
        "UNSUPPORTED": sum(
            clean_text(row.get("audit_status")) == "UNSUPPORTED"
            for row in adjudication
        ),
    }
    citation_targets = {
        "total": int(expected["citation_rows_total"]),
        "SUPPORTED": int(expected["citation_rows_supported"]),
        "PARTIAL": int(expected["citation_rows_partial"]),
        "UNSUPPORTED": int(expected["citation_rows_unsupported"]),
    }
    if citation_stats != citation_targets:
        raise Phase8ExportError("Citation adjudication counts changed")

    scoring = _sheet_rows(workbook["Scoring Criteria"])
    total_available = sum(_as_int(row.get("points_available")) for row in scoring)
    numeric_awards = [
        float(row["points_awarded"])
        for row in scoring
        if row.get("points_awarded") not in (None, "")
    ]
    if total_available != int(expected["scoring_points_available"]):
        raise Phase8ExportError("Tender scoring model no longer totals 100")
    if numeric_awards and abs(sum(numeric_awards)) > 1e-9:
        raise Phase8ExportError("Phase 8 must not award tender points")

    records = []
    for row in sorted(shortlist, key=lambda item: _as_int(item.get("audit_rank"))):
        reference_id = clean_text(row["reference_id"])
        role = clean_text(row["verified_primary_role"])
        copy = MISSION_COPY[reference_id]
        records.append(
            {
                "rank": _as_int(row.get("audit_rank")),
                "reference_id": reference_id,
                "client": clean_text(row.get("client")),
                "country_code": clean_text(row.get("country_code")),
                "country": (
                    "Tunisie"
                    if clean_text(row.get("country_code")) == "TN"
                    else "Côte d'Ivoire"
                    if clean_text(row.get("country_code")) == "CI"
                    else clean_text(row.get("country_code"))
                ),
                "subregion": clean_text(row.get("subregion")),
                "sector_code": clean_text(row.get("sector")),
                "sector": SECTOR_LABELS_FR.get(
                    clean_text(row.get("sector")), clean_text(row.get("sector"))
                ),
                "verified_primary_role": role,
                "verified_role_label": ROLE_LABELS_FR[role],
                "mission_title": copy["title"],
                "verified_summary": copy["summary"],
                "mission_period_verified": clean_text(
                    row.get("mission_period_verified")
                ),
                "evidence_type": clean_text(row.get("evidence_type")),
                "evidence_type_label": EVIDENCE_LABELS_FR.get(
                    clean_text(row.get("evidence_type")),
                    clean_text(row.get("evidence_type")).replace("_", " "),
                ),
                "engine_score_not_tender_points": float(
                    row.get("engine_score_not_tender_points") or 0.0
                ),
                "source_file_name": clean_text(row.get("source_file_name")),
                "source_url": clean_text(row.get("source_url")),
                "provenance": clean_text(row.get("provenance")),
                "citation_support": "SUPPORTED",
                "approval_status": approval["approval_status"],
                "audit_notes": clean_text(row.get("audit_notes")),
                "covered_requirements": _requirement_rows_for_reference(
                    row, coverage
                ),
            }
        )

    return {
        "label": LABEL,
        "status": STATUS,
        "opportunity_id": config["inputs"]["opportunity_id"],
        "approval": approval,
        "summary": summary,
        "records": records,
        "eligibility": eligibility,
        "coverage": coverage,
        "citation_stats": citation_stats,
        "scoring": scoring,
        "metrics": {
            "selected_references": len(records),
            "eligibility_passed": len(passed),
            "eligibility_total": len(eligibility),
            "must_covered": len(must_covered),
            "must_total": len(must),
            "should_covered": len(should_covered),
            "should_total": len(should),
            "explicit_gaps": gaps,
            "points_available": total_available,
            "points_awarded": 0,
        },
    }


def _set_run_font(
    run,
    *,
    name: str = "Arial",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_cell_margins(
    cell,
    *,
    top: int = 80,
    start: int = 110,
    bottom: int = 80,
    end: int = 110,
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def _set_table_geometry(table, widths_cm: list[float], indent_dxa: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    widths_dxa = [int(round(width * 567.0)) for width in widths_cm]
    total = sum(widths_dxa)
    table.autofit = False
    properties = table._tbl.tblPr

    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.insert(0, table_width)
    table_width.set(qn("w:w"), str(total))
    table_width.set(qn("w:type"), "dxa")

    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), str(indent_dxa))
    indent.set(qn("w:type"), "dxa")

    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = width
            tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")


def _set_repeat_header(row) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = row._tr.get_or_add_trPr()
    marker = properties.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        properties.append(marker)
    marker.set(qn("w:val"), "true")


def _prevent_row_split(row) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = row._tr.get_or_add_trPr()
    marker = properties.find(qn("w:cantSplit"))
    if marker is None:
        marker = OxmlElement("w:cantSplit")
        properties.append(marker)


def _format_cell(
    cell,
    *,
    size: float = 8.3,
    bold: bool = False,
    color: str = "1F2937",
    alignment=None,
    background: str | None = None,
) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Pt

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)
    if background:
        _set_cell_shading(cell, background)
    for paragraph in cell.paragraphs:
        if alignment is not None:
            paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            _set_run_font(
                run,
                size=size,
                color=color,
                bold=bold,
            )


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([fonts, color, underline])
    value = OxmlElement("w:t")
    value.text = text
    run.extend([properties, value])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_page_field(paragraph, field: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    run.extend([begin, instruction, separate, text, end])
    paragraph._p.append(run)


def _shade_paragraph(paragraph, fill: str, border: str | None = None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = paragraph._p.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    if border:
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), border)
        borders.append(bottom)
        properties.append(borders)


def _configure_styles(document) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string("1F2937")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    tokens = {
        "Heading 1": (16, "B8162E", 16, 8),
        "Heading 2": (13, "B8162E", 12, 6),
        "Heading 3": (11, "374151", 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Source Citation" not in styles:
        styles.add_style("Source Citation", WD_STYLE_TYPE.PARAGRAPH)
    source = styles["Source Citation"]
    source.font.name = "Arial"
    source._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    source._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    source.font.size = Pt(8.5)
    source.font.color.rgb = RGBColor.from_string("4B5563")
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)


def _configure_section(section, *, landscape: bool) -> None:
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm

    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
    else:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)


def _add_header_footer(section) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    section.header.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("SYNTHETIC_TEST_ONLY | BRID | Phase 8")
    _set_run_font(run, size=8, color="6B7280", bold=True)

    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run(
        "DRAFT INTERNE - SYNTHETIC_TEST_ONLY - validation métier et sécurité requise | Page "
    )
    _set_run_font(run, size=7.5, color="6B7280")
    _add_page_field(footer, "PAGE")
    run = footer.add_run(" / ")
    _set_run_font(run, size=7.5, color="6B7280")
    _add_page_field(footer, "NUMPAGES")


def _add_status_banner(document, text: str, *, width_cm: float) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(4)
    paragraph.paragraph_format.right_indent = Pt(4)
    _shade_paragraph(paragraph, "B8162E")
    run = paragraph.add_run(text)
    _set_run_font(run, size=9, color="FFFFFF", bold=True)


def _add_kpi_strip(document, metrics: dict) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    values = [
        ("Références", str(metrics["selected_references"])),
        (
            "Éligibilité",
            f"{metrics['eligibility_passed']}/{metrics['eligibility_total']}",
        ),
        ("MUST", f"{metrics['must_covered']}/{metrics['must_total']}"),
        ("SHOULD", f"{metrics['should_covered']}/{metrics['should_total']}"),
        ("Points attribués", f"{metrics['points_awarded']}/100"),
    ]
    table = document.add_table(rows=1, cols=len(values))
    for index, (label, value) in enumerate(values):
        cell = table.cell(0, index)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(value)
        _set_run_font(run, size=14, color="B8162E", bold=True)
        run = paragraph.add_run(f"\n{label}")
        _set_run_font(run, size=7.8, color="4B5563", bold=True)
        _set_cell_margins(cell, top=110, bottom=110, start=90, end=90)
        _set_cell_shading(cell, "F9FAFB")
    _set_table_geometry(table, [5.1, 5.1, 5.1, 5.1, 5.1], indent_dxa=0)
    _set_repeat_header(table.rows[0])


def _add_heading(document, text: str, level: int = 1):
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def _add_summary_table(document, records: list[dict]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    headers = (
        "#",
        "Mission vérifiée",
        "Client",
        "Pays",
        "Période",
        "Rôle audité",
        "Score moteur*",
        "Preuve",
        "Décision",
    )
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.cell(0, index).text = value
        _format_cell(
            table.cell(0, index),
            size=7.4,
            bold=True,
            color="FFFFFF",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            background="F9425B",
        )
    _set_repeat_header(table.rows[0])
    for record in records:
        cells = table.add_row().cells
        values = (
            str(record["rank"]),
            record["mission_title"],
            record["client"],
            record["country"],
            record["mission_period_verified"],
            record["verified_role_label"],
            f"{record['engine_score_not_tender_points']:.3f}",
            record["evidence_type_label"],
            "APPROUVÉ",
        )
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            _format_cell(
                cell,
                size=7.2,
                bold=index == 0,
                alignment=(
                    WD_ALIGN_PARAGRAPH.CENTER
                    if index in {0, 3, 4, 6, 7, 8}
                    else WD_ALIGN_PARAGRAPH.LEFT
                ),
                background="FFF5F6" if record["rank"] % 2 == 0 else None,
            )
        _prevent_row_split(table.rows[-1])
    _set_table_geometry(
        table,
        [0.8, 4.8, 3.8, 1.6, 2.6, 3.4, 2.0, 3.2, 2.0],
        indent_dxa=0,
    )


def _add_eligibility_table(document, rows: list[dict]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    headers = ("Règle", "Requis", "Observé", "Statut", "Références qualifiantes", "Base d'audit")
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.cell(0, index).text = value
        _format_cell(
            table.cell(0, index),
            size=8,
            bold=True,
            color="FFFFFF",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            background="F9425B",
        )
    _set_repeat_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        values = (
            clean_text(row.get("rule_code")),
            str(_as_int(row.get("required_count"))),
            str(_as_int(row.get("observed_count"))),
            clean_text(row.get("audit_status")),
            clean_text(row.get("qualifying_clients")),
            ELIGIBILITY_BASIS_FR.get(
                clean_text(row.get("rule_code")), clean_text(row.get("audit_basis"))
            ),
        )
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            _format_cell(
                cell,
                size=7.5,
                bold=index == 3,
                color="166534" if index == 3 else "1F2937",
                alignment=(
                    WD_ALIGN_PARAGRAPH.CENTER if index in {1, 2, 3} else WD_ALIGN_PARAGRAPH.LEFT
                ),
                background="F0FDF4" if index == 3 else None,
            )
        _prevent_row_split(table.rows[-1])
    _set_table_geometry(table, [2.6, 1.4, 1.5, 1.7, 4.9, 5.3], indent_dxa=0)


def _add_coverage_table(document, rows: list[dict]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    headers = ("Code", "Classe", "Statut", "Références", "Exigence auditée")
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.cell(0, index).text = value
        _format_cell(
            table.cell(0, index),
            size=8,
            bold=True,
            color="FFFFFF",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            background="F9425B",
        )
    _set_repeat_header(table.rows[0])
    for row in rows:
        status = clean_text(row.get("audit_status"))
        cells = table.add_row().cells
        values = (
            clean_text(row.get("code")),
            clean_text(row.get("classification")),
            status,
            str(_as_int(row.get("supporting_reference_count"))),
            clean_text(row.get("requirement_text")),
        )
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            _format_cell(
                cell,
                size=7.8,
                bold=index == 2,
                color=(
                    "166534" if index == 2 and status == "COVERED"
                    else "9A3412" if index == 2 else "1F2937"
                ),
                alignment=(
                    WD_ALIGN_PARAGRAPH.CENTER if index in {0, 1, 2, 3} else WD_ALIGN_PARAGRAPH.LEFT
                ),
                background=(
                    "F0FDF4" if index == 2 and status == "COVERED"
                    else "FFF7ED" if index == 2 else None
                ),
            )
        _prevent_row_split(table.rows[-1])
    _set_table_geometry(table, [2.5, 1.7, 2.0, 1.8, 9.4], indent_dxa=0)


def _add_scoring_table(document, rows: list[dict]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    headers = ("Code", "Critère", "Points disponibles", "Route d'évaluation", "Points attribués")
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.cell(0, index).text = value
        _format_cell(
            table.cell(0, index),
            size=8,
            bold=True,
            color="FFFFFF",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            background="374151",
        )
    _set_repeat_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        values = (
            clean_text(row.get("code")),
            clean_text(row.get("criterion")),
            str(_as_int(row.get("points_available"))),
            clean_text(row.get("evaluation_route")),
            "0 - non évalué",
        )
        for index, (cell, value) in enumerate(zip(cells, values)):
            cell.text = value
            _format_cell(
                cell,
                size=7.8,
                bold=index == 4,
                color="9A3412" if index == 4 else "1F2937",
                alignment=(
                    WD_ALIGN_PARAGRAPH.CENTER if index in {0, 2, 3, 4} else WD_ALIGN_PARAGRAPH.LEFT
                ),
                background="FFF7ED" if index == 4 else None,
            )
        _prevent_row_split(table.rows[-1])
    _set_table_geometry(table, [2.8, 6.5, 2.3, 3.3, 2.5], indent_dxa=0)


def _add_detail_page(document, record: dict, position: int) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document.add_page_break()
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(3)
    run = heading.add_run(f"Référence N°{position}")
    _set_run_font(run, size=10, color="B8162E", bold=True)
    title = document.add_paragraph(style="Heading 2")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(record["mission_title"])
    _set_run_font(run, size=16, color="1F2937", bold=True)

    fields = [
        ("Client", record["client"]),
        ("Pays / région", f"{record['country']} - {record['subregion']}"),
        ("Secteur", record["sector"]),
        ("Période vérifiée", record["mission_period_verified"]),
        ("Rôle principal audité", record["verified_role_label"]),
        ("Type de preuve", record["evidence_type_label"]),
        ("Statut de preuve", record["citation_support"]),
        (
            "Score moteur",
            f"{record['engine_score_not_tender_points']:.3f} - classement technique, pas des points d'appel d'offres",
        ),
    ]
    table = document.add_table(rows=len(fields) + 1, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Champ"
    table.cell(0, 1).text = "Information auditée"
    for cell in table.rows[0].cells:
        _format_cell(
            cell,
            size=8,
            bold=True,
            color="FFFFFF",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            background="F9425B",
        )
    _set_repeat_header(table.rows[0])
    for row_index, (label, value) in enumerate(fields, start=1):
        table.cell(row_index, 0).text = label
        table.cell(row_index, 1).text = value
        _format_cell(
            table.cell(row_index, 0),
            size=8.5,
            bold=True,
            background="F5A0AE",
        )
        _format_cell(table.cell(row_index, 1), size=8.5)
        _prevent_row_split(table.rows[row_index])
    _set_table_geometry(table, [4.2, 13.2], indent_dxa=0)

    _add_heading(document, "Portée vérifiée", level=3)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(8)
    paragraph.paragraph_format.right_indent = Pt(8)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    _shade_paragraph(paragraph, "F9FAFB", "D1D5DB")
    run = paragraph.add_run(record["verified_summary"])
    _set_run_font(run, size=9.5, color="1F2937")

    _add_heading(document, "Exigences couvertes par cette référence", level=3)
    requirements = record["covered_requirements"]
    coverage_table = document.add_table(rows=1, cols=3)
    coverage_table.style = "Table Grid"
    for index, value in enumerate(("Code", "Classe", "Exigence")):
        coverage_table.cell(0, index).text = value
        _format_cell(
            coverage_table.cell(0, index),
            size=8,
            bold=True,
            color="FFFFFF",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            background="F9425B",
        )
    _set_repeat_header(coverage_table.rows[0])
    for requirement in requirements:
        cells = coverage_table.add_row().cells
        cells[0].text = requirement["code"]
        cells[1].text = requirement["classification"]
        cells[2].text = requirement["requirement_text"]
        for index, cell in enumerate(cells):
            _format_cell(
                cell,
                size=8,
                alignment=(
                    WD_ALIGN_PARAGRAPH.CENTER if index < 2 else WD_ALIGN_PARAGRAPH.LEFT
                ),
            )
        _prevent_row_split(coverage_table.rows[-1])
    _set_table_geometry(coverage_table, [2.7, 2.0, 12.7], indent_dxa=0)

    _add_heading(document, "Justificatif audité", level=3)
    source = document.add_paragraph(style="Source Citation")
    run = source.add_run(f"{record['source_file_name']} - ")
    _set_run_font(run, size=8.5, color="4B5563", bold=True)
    _add_hyperlink(source, "Ouvrir la source exacte (page 1)", record["source_url"])
    source.add_run(
        f"\nProvenance : {record['provenance']} | Support : {record['citation_support']}"
    )

    boundary = document.add_paragraph()
    boundary.paragraph_format.space_before = Pt(5)
    boundary.paragraph_format.space_after = Pt(0)
    _shade_paragraph(boundary, "FFF7ED", "FDBA74")
    run = boundary.add_run(
        "Limite d'utilisation : cette fiche reprend uniquement la portée confirmée "
        "par l'audit. Elle ne constitue ni une note commerciale, ni une attribution "
        "automatique de points."
    )
    _set_run_font(run, size=8.2, color="9A3412", bold=True)


def create_reference_docx(path: Path, data: dict, template: dict) -> None:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    document.core_properties.title = "BRID - Dossier de références audité"
    document.core_properties.subject = "Phase 8 - SYNTHETIC_TEST_ONLY"
    document.core_properties.author = "Devoteam Reference AI"
    document.core_properties.comments = (
        "Generated deterministically from the approved Phase 7 evidence audit."
    )
    _configure_styles(document)

    first = document.sections[0]
    _configure_section(first, landscape=True)
    _add_header_footer(first)
    _add_status_banner(
        document,
        "SYNTHETIC_TEST_ONLY - LISTE AUDITÉE APPROUVÉE - AUCUNE LIVRAISON CLIENT AUTOMATIQUE",
        width_cm=27.3,
    )

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Dossier de références audité - Cas BRID")
    _set_run_font(run, size=23, color="1F2937", bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run(
        "Phase 8 | Export déterministe fondé sur l'audit de preuves | "
        f"Opportunité {data['opportunity_id']}"
    )
    _set_run_font(run, size=10, color="6B7280", italic=True)

    _add_kpi_strip(document, data["metrics"])

    gap = document.add_paragraph()
    gap.paragraph_format.space_before = Pt(7)
    gap.paragraph_format.space_after = Pt(7)
    _shade_paragraph(gap, "FFF7ED", "FDBA74")
    run = gap.add_run(
        "Écart explicite non bloquant : SCOPE-CLOUD. "
        "La liste couvre 8/8 exigences MUST et 2/3 exigences SHOULD. "
        "Aucun point d'appel d'offres n'a été attribué."
    )
    _set_run_font(run, size=9, color="9A3412", bold=True)

    _add_heading(document, "I. Nos principales références", level=1)
    _add_summary_table(document, data["records"])
    note = document.add_paragraph(style="Source Citation")
    note.paragraph_format.space_before = Pt(4)
    run = note.add_run(
        "* Le score moteur sert uniquement au classement technique. "
        "Il ne correspond pas aux 100 points du modèle d'évaluation."
    )
    _set_run_font(run, size=8, color="6B7280", italic=True)

    controls = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(controls, landscape=False)
    _add_header_footer(controls)

    _add_status_banner(
        document,
        "CONTRÔLE HUMAIN ENREGISTRÉ - EXPORT SYNTHÉTIQUE UNIQUEMENT",
        width_cm=17.4,
    )
    _add_heading(document, "II. Contrôles d'éligibilité et de couverture", level=1)
    approval = document.add_paragraph()
    approval.paragraph_format.space_after = Pt(8)
    run = approval.add_run(
        "Décision : liste auditée approuvée. "
        f"Audit de citations : {data['citation_stats']['SUPPORTED']} supportées, "
        f"{data['citation_stats']['PARTIAL']} partielles et "
        f"{data['citation_stats']['UNSUPPORTED']} non supportées sur "
        f"{data['citation_stats']['total']} lignes automatiques."
    )
    _set_run_font(run, size=9.5, color="1F2937")

    _add_heading(document, "Éligibilité du portefeuille", level=2)
    _add_eligibility_table(document, data["eligibility"])

    _add_heading(document, "Couverture des exigences", level=2)
    _add_coverage_table(document, data["coverage"])

    _add_heading(document, "Modèle de notation - non exécuté", level=2)
    _add_scoring_table(document, data["scoring"])
    scoring_note = document.add_paragraph()
    _shade_paragraph(scoring_note, "FFF7ED", "FDBA74")
    run = scoring_note.add_run(
        "La grille totalise 100 points, mais Phase 8 en attribue 0. "
        "La notation exige une évaluation humaine séparée, notamment pour "
        "l'équipe et la méthodologie."
    )
    _set_run_font(run, size=8.8, color="9A3412", bold=True)

    _add_heading(document, "Traçabilité du gabarit", level=2)
    trace = document.add_paragraph()
    run = trace.add_run(
        "Le gabarit de référence a été vérifié avant génération : "
        f"SHA-256 {template['sha256']}; {template['tables']} tableaux; "
        f"{template['reference_slots']} emplacements de références."
    )
    _set_run_font(run, size=8.8, color="4B5563")

    _add_heading(document, "III. Fiches détaillées", level=1)
    intro = document.add_paragraph(
        "Les fiches suivantes reprennent uniquement les rôles et la portée "
        "confirmés par l'audit de preuves."
    )
    intro.paragraph_format.space_after = Pt(0)
    for position, record in enumerate(data["records"], start=1):
        _add_detail_page(document, record, position)

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def _build_report(data: dict, template: dict) -> str:
    metrics = data["metrics"]
    lines = [
        "# Phase 8 BRID - audited reference export",
        "",
        f"**Label:** `{LABEL}`  ",
        f"**Status:** `{STATUS}`  ",
        f"**Opportunity:** `{data['opportunity_id']}`  ",
        f"**Approval:** `{data['approval']['approval_status']}`  ",
        "",
        "## Outcome",
        "",
        f"- {metrics['selected_references']} audited references exported.",
        (
            f"- Eligibility: {metrics['eligibility_passed']}/"
            f"{metrics['eligibility_total']} rules passed."
        ),
        f"- MUST coverage: {metrics['must_covered']}/{metrics['must_total']}.",
        f"- SHOULD coverage: {metrics['should_covered']}/{metrics['should_total']}.",
        "- Explicit non-blocking gap: `SCOPE-CLOUD`.",
        "- Tender points awarded: 0/100.",
        "- Phase 5.1 expert relevance evaluation remains pending.",
        "",
        "## Audited references",
        "",
    ]
    for record in data["records"]:
        lines.append(
            f"{record['rank']}. **{record['client']}** - "
            f"{record['mission_title']} "
            f"([{record['source_file_name']}]({record['source_url']}))."
        )
    lines.extend(
        [
            "",
            "## Controls",
            "",
            (
                f"- Citation audit: {data['citation_stats']['SUPPORTED']} supported, "
                f"{data['citation_stats']['PARTIAL']} partial, "
                f"{data['citation_stats']['UNSUPPORTED']} unsupported."
            ),
            (
                "- Every exported shortlist row is marked `SUPPORTED` and has an "
                "exact Drive/page source."
            ),
            (
                f"- Reference template verified: `{template['sha256']}` "
                f"({template['tables']} tables)."
            ),
            "- No LLM, translation API, automatic tender scoring, or client delivery was used.",
            "- The original Phase 7 audit workbook was not modified.",
            "",
            "## Decision boundary",
            "",
            (
                "This release completes the controlled synthetic Phase 8 export. "
                "It does not approve a real tender response, replace Phase 5.1 expert "
                "judgments, or authorize production/client use."
            ),
            "",
        ]
    )
    return "\n".join(lines)


def build_pre_pdf_outputs(
    *,
    audit_workbook: Path,
    reference_template: Path,
    approval_path: Path,
    config_path: Path,
    output_dir: Path,
) -> tuple[dict, dict, dict]:
    config = load_phase8_config(config_path)
    audit = verify_audit_workbook(audit_workbook, config)
    template = verify_reference_template(reference_template, config)
    approval = load_approval(approval_path, config, audit["sha256"])
    data = load_audited_case(audit_workbook, config, approval)
    output_dir.mkdir(parents=True, exist_ok=True)
    outputs = config["output"]

    create_reference_docx(output_dir / outputs["docx_name"], data, template)
    shutil.copy2(approval_path, output_dir / outputs["approval_name"])
    _atomic_json(output_dir / outputs["data_name"], data)
    _atomic_text(output_dir / outputs["report_name"], _build_report(data, template))
    return config, data, template


def finalize_run(
    *,
    config_path: Path,
    audit_workbook: Path,
    reference_template: Path,
    output_dir: Path,
    visual_qa_status: str,
) -> dict:
    from docx import Document
    from pypdf import PdfReader

    config = load_phase8_config(config_path)
    outputs = config["output"]
    docx_path = output_dir / outputs["docx_name"]
    pdf_path = output_dir / outputs["pdf_name"]
    data_path = output_dir / outputs["data_name"]
    approval_path = output_dir / outputs["approval_name"]
    report_path = output_dir / outputs["report_name"]
    for path in (docx_path, pdf_path, data_path, approval_path, report_path):
        if not path.exists():
            raise FileNotFoundError(path)
    if not pdf_path.read_bytes().startswith(b"%PDF"):
        raise AssertionError("Phase 8 PDF signature is invalid")

    data = json.loads(data_path.read_text(encoding="utf-8"))
    document = Document(docx_path)
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    combined = body_text + "\n" + table_text
    for record in data["records"]:
        if record["client"] not in combined or record["mission_title"] not in combined:
            raise AssertionError(f"Reference missing from DOCX: {record['client']}")
    if "SCOPE-CLOUD" not in combined or "0/100" not in combined:
        raise AssertionError("Required Phase 8 caveats are missing from the DOCX")

    relationships = []
    for relationship in document.part.rels.values():
        if relationship.reltype.endswith("/hyperlink"):
            relationships.append(str(relationship.target_ref))
    source_urls = {record["source_url"] for record in data["records"]}
    if not source_urls.issubset(set(relationships)):
        raise AssertionError("Not every audited source is linked in the DOCX")

    pdf_reader = PdfReader(str(pdf_path))
    if len(pdf_reader.pages) < 12:
        raise AssertionError("Phase 8 PDF has fewer pages than expected")

    audit_sha = sha256_file(audit_workbook)
    template_sha = sha256_file(reference_template)
    manifest = {
        "schema_version": 1,
        "phase": 8,
        "pipeline_version": config["pipeline_version"],
        "completed_at_utc": utc_now(),
        "status": STATUS,
        "label": LABEL,
        "opportunity_id": data["opportunity_id"],
        "approval_status": data["approval"]["approval_status"],
        "approval_scope": data["approval"]["approval_scope"],
        "approver_identity_recorded": False,
        "source_audit_workbook_sha256": audit_sha,
        "source_reference_template_sha256": template_sha,
        "source_audit_workbook_unchanged": (
            audit_sha == config["inputs"]["audit_workbook_sha256"]
        ),
        "selected_references": data["metrics"]["selected_references"],
        "eligibility_rules_passed": data["metrics"]["eligibility_passed"],
        "eligibility_rules_total": data["metrics"]["eligibility_total"],
        "must_covered": data["metrics"]["must_covered"],
        "must_total": data["metrics"]["must_total"],
        "should_covered": data["metrics"]["should_covered"],
        "should_total": data["metrics"]["should_total"],
        "explicit_gaps": data["metrics"]["explicit_gaps"],
        "citation_adjudication": data["citation_stats"],
        "scoring_points_available": data["metrics"]["points_available"],
        "scoring_points_awarded": 0,
        "phase5_1_expert_evaluation": "PENDING",
        "external_llm_calls": 0,
        "external_translation_calls": 0,
        "automatic_tender_scoring": False,
        "automatic_client_delivery": False,
        "production_promotion_allowed": False,
        "source_documents_copied": False,
        "source_page_images_embedded": False,
        "template_structure_verified": True,
        "visual_qa_status": visual_qa_status,
        "docx_hyperlink_count": len(relationships),
        "pdf_page_count": len(pdf_reader.pages),
    }
    manifest_path = output_dir / outputs["manifest_name"]
    _atomic_json(manifest_path, manifest)

    included = sorted(
        (
            docx_path,
            pdf_path,
            data_path,
            approval_path,
            report_path,
            manifest_path,
        ),
        key=lambda item: item.name,
    )
    sums = "".join(f"{sha256_file(path)}  {path.name}\n" for path in included)
    sums_path = output_dir / outputs["checksum_name"]
    _atomic_text(sums_path, sums)
    success = {
        "status": "COMPLETE_REPRODUCIBLE_AUDITED_REFERENCE_DOSSIER",
        "created_at_utc": utc_now(),
        "pipeline_version": config["pipeline_version"],
        "opportunity_id": data["opportunity_id"],
        "manifest_sha256": sha256_file(manifest_path),
        "sha256sums_sha256": sha256_file(sums_path),
    }
    _atomic_json(output_dir / outputs["success_name"], success)
    verify_run(output_dir=output_dir, config_path=config_path)
    return manifest


def verify_run(*, output_dir: Path, config_path: Path) -> dict:
    config = load_phase8_config(config_path)
    outputs = config["output"]
    manifest_path = output_dir / outputs["manifest_name"]
    sums_path = output_dir / outputs["checksum_name"]
    success_path = output_dir / outputs["success_name"]
    for path in (manifest_path, sums_path, success_path):
        if not path.exists():
            raise FileNotFoundError(path)
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    success = json.loads(success_path.read_text(encoding="utf-8"))
    if success["manifest_sha256"] != sha256_file(manifest_path):
        raise AssertionError("Phase 8 manifest hash mismatch")
    if success["sha256sums_sha256"] != sha256_file(sums_path):
        raise AssertionError("Phase 8 checksum file hash mismatch")
    for line in sums_path.read_text(encoding="utf-8").splitlines():
        expected, name = line.split("  ", 1)
        if sha256_file(output_dir / name) != expected:
            raise AssertionError(f"Phase 8 output changed: {name}")
    if manifest.get("status") != STATUS or manifest.get("label") != LABEL:
        raise AssertionError("Unexpected Phase 8 status or label")
    if manifest.get("approval_status") != "APPROVED":
        raise AssertionError("Phase 8 approval gate failed")
    if int(manifest.get("selected_references", 0)) != 10:
        raise AssertionError("Phase 8 reference count changed")
    if (
        int(manifest.get("eligibility_rules_passed", 0)) != 7
        or int(manifest.get("must_covered", 0)) != 8
        or int(manifest.get("should_covered", 0)) != 2
    ):
        raise AssertionError("Phase 8 eligibility or coverage gate failed")
    if manifest.get("explicit_gaps") != ["SCOPE-CLOUD"]:
        raise AssertionError("Phase 8 explicit gap changed")
    if int(manifest.get("scoring_points_awarded", -1)) != 0:
        raise AssertionError("Phase 8 awarded tender points")
    if any(
        bool(manifest.get(field))
        for field in (
            "automatic_tender_scoring",
            "automatic_client_delivery",
            "production_promotion_allowed",
            "source_documents_copied",
            "source_page_images_embedded",
        )
    ):
        raise AssertionError("Phase 8 safety boundary failed")
    if manifest.get("visual_qa_status") != "PASSED":
        raise AssertionError("Phase 8 visual QA gate failed")
    return {"manifest": manifest, "success": success}
