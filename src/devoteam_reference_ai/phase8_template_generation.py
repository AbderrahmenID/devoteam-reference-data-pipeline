from __future__ import annotations

import hashlib
import json
import math
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

import pandas as pd
import yaml

from .phase6_opportunity import sha256_file
from .phase7_recommendations import verify_phase7_run


YEAR_RE = re.compile(r"\b(?:19|20)\d{2}\b")
STOPWORDS = {
    "avec", "dans", "pour", "sans", "sous", "plus", "moins", "entre", "ainsi", "afin",
    "cette", "leurs", "notre", "votre", "dont", "tout", "tous", "toute", "toutes", "être",
    "avoir", "faire", "mise", "place", "projet", "mission", "phase", "article", "contrat",
    "elaboration", "élaboration",
    "devoteam", "prestataire", "client", "banque", "the", "and", "from", "with", "this",
}


class TemplateGenerationError(RuntimeError):
    """Raised when shortlist review, template controls, or evidence are incomplete."""


def utc_now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _atomic_text(path: Path, value: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(value, encoding="utf-8")
    os.replace(temporary, path)


def _atomic_json(path: Path, value: Any) -> None:
    _atomic_text(path, json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True) + "\n")


def clean_text(value: Any) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return ""
    return " ".join(str(value).replace("\u00a0", " ").split())


def load_phase8_config(path: Path) -> dict:
    config = yaml.safe_load(path.read_text(encoding="utf-8"))
    if int(config.get("phase", 0)) != 8:
        raise ValueError("Expected Phase 8 configuration")
    generation = config.get("generation", {})
    blocked = (
        "external_llm_enabled",
        "external_translation_enabled",
        "automatic_client_delivery",
        "copy_source_documents_into_output",
        "embed_source_page_images",
    )
    if any(bool(generation.get(field)) for field in blocked):
        raise ValueError("Phase 8 generation/security defaults were weakened")
    security = config.get("security", {})
    if any(
        bool(security.get(field))
        for field in ("raw_opportunity_logging_allowed", "source_mutation_allowed", "hidden_content_generation_allowed")
    ):
        raise ValueError("Phase 8 security defaults were weakened")
    if generation.get("mode") != "deterministic_template_derived":
        raise ValueError("The signed Phase 8 package is deterministic and template-derived")
    if int(generation.get("maximum_references", 0)) not in range(1, 21):
        raise ValueError("maximum_references must be between 1 and 20")
    return config


def verify_reference_template(path: Path, config: dict) -> dict:
    from docx import Document

    if not path.exists():
        raise FileNotFoundError(path)
    expected = config["input"]
    if path.stat().st_size != int(expected["expected_template_size_bytes"]):
        raise AssertionError("Reference template size changed")
    digest = sha256_file(path)
    if digest != expected["expected_template_sha256"]:
        raise AssertionError("Reference template hash changed")
    document = Document(path)
    if len(document.tables) < 18:
        raise AssertionError("Reference template detail-table structure changed")
    summary = document.tables[0]
    if len(summary.rows) < 18 or len(summary.columns) < 10:
        raise AssertionError("Reference template summary-table structure changed")
    summary_text = " ".join(cell.text for cell in summary.rows[0].cells + summary.rows[1].cells)
    for label in ("Intitulé du projet", "Client", "Pays", "Période", "Thématiques clés"):
        if label not in summary_text:
            raise AssertionError(f"Reference template label missing: {label}")
    detail_labels = {clean_text(row.cells[0].text) for row in document.tables[1].rows}
    for label in ("Nom de la mission", "Pays", "Nom de l’Autorité Contractante", "Date de démarrage"):
        if label not in detail_labels:
            raise AssertionError(f"Reference detail label missing: {label}")
    return {
        "sha256": digest,
        "size_bytes": path.stat().st_size,
        "tables": len(document.tables),
        "sections": len(document.sections),
        "reference_slots": len(document.tables) - 1,
    }


def _sheet_rows(sheet) -> list[dict]:
    values = list(sheet.iter_rows(values_only=True))
    if not values:
        return []
    headers = [clean_text(value) for value in values[0]]
    return [dict(zip(headers, row)) for row in values[1:] if any(value not in (None, "") for value in row)]


def load_selected_shortlist(phase7_root: Path, phase7_config: dict, phase8_config: dict) -> dict:
    verification = verify_phase7_run(phase7_root, phase7_config)
    manifest = verification["manifest"]
    recommendations = pd.read_parquet(phase7_root / "recommendations.parquet")
    evidence = pd.read_parquet(phase7_root / "evidence_matrix.parquet")
    coverage = pd.read_csv(phase7_root / "requirement_coverage.csv", encoding="utf-8-sig")
    if recommendations.empty:
        raise TemplateGenerationError("Phase 7 produced no recommendations")
    if manifest["status"] == "TECHNICAL_PASS_SAMPLE_ONLY":
        if not phase8_config["review"]["synthetic_sample_auto_selection_allowed"]:
            raise TemplateGenerationError("Synthetic sample selection is disabled")
        selected = recommendations.head(int(phase8_config["review"]["synthetic_sample_top_n"])).copy()
        selection_mode = "SYNTHETIC_DEVELOPMENT_FIXTURE_TOP_N"
    else:
        import openpyxl

        workbook = openpyxl.load_workbook(
            phase7_root / "RECOMMENDATION_REVIEW.xlsx", data_only=False, read_only=True
        )
        reviews = _sheet_rows(workbook["Ranked References"])
        decisions = {clean_text(row.get("reference_id")): clean_text(row.get("reviewer_decision")).upper() for row in reviews}
        recommendation_ids = recommendations["reference_id"].astype(str).tolist()
        invalid = {decision for decision in decisions.values() if decision not in {"PENDING", "SHORTLIST", "REJECT"}}
        if invalid:
            raise TemplateGenerationError(f"Unknown shortlist decisions: {sorted(invalid)}")
        if any(decisions.get(reference_id, "PENDING") == "PENDING" for reference_id in recommendation_ids):
            raise TemplateGenerationError("Every real recommendation must be reviewed before generation")
        selected_ids = [reference_id for reference_id in recommendation_ids if decisions[reference_id] == "SHORTLIST"]
        if not selected_ids:
            raise TemplateGenerationError("At least one reference must be marked SHORTLIST")
        selected = recommendations.loc[recommendations["reference_id"].astype(str).isin(selected_ids)].copy()
        selection_mode = "HUMAN_APPROVED_SHORTLIST"
    maximum = int(phase8_config["generation"]["maximum_references"])
    selected = selected.sort_values("final_rank").head(maximum).reset_index(drop=True)
    selected_ids = set(selected["reference_id"].astype(str))
    selected_evidence = evidence.loc[evidence["reference_id"].astype(str).isin(selected_ids)].copy()
    selected_coverage = coverage.loc[coverage["reference_id"].astype(str).isin(selected_ids)].copy()
    evidence_ids = set(selected_evidence["reference_id"].astype(str))
    if evidence_ids != selected_ids:
        raise TemplateGenerationError("Every selected reference must have evidence")
    citation_fields = ["citation_uri", "citation_label", "evidence_excerpt", "source_sha256", "chunk_text_sha256"]
    if selected_evidence[citation_fields].isna().any().any() or selected_evidence[citation_fields].eq("").any().any():
        raise TemplateGenerationError("Selected evidence has incomplete provenance")
    return {
        "phase7_manifest": manifest,
        "phase7_manifest_sha256": sha256_file(phase7_root / "PHASE_7_MANIFEST.json"),
        "recommendations": selected,
        "evidence": selected_evidence,
        "coverage": selected_coverage,
        "selection_mode": selection_mode,
    }


def project_year_bounds(value: Any) -> tuple[str, str]:
    years = YEAR_RE.findall(str(value or ""))
    if not years:
        return "À confirmer", "À confirmer"
    return years[0], years[-1]


def mission_title(reference: dict) -> str:
    service = clean_text(reference.get("service_nature"))
    first = re.split(
        r"\s*[•▪]\s*|\s+(?:PHASE|Phase|ETAPE|Etape|ÉTAPE|Étape)\s+\d+\b|(?<=[.!?;])\s+",
        service,
        maxsplit=1,
    )[0].strip(" -•")
    if 20 <= len(first) <= 180:
        return first
    offering = clean_text(reference.get("offering")) or "Mission de conseil"
    client = clean_text(reference.get("client")) or "Client à confirmer"
    return f"{offering} — {client}"


def lexical_grounding_score(context: Any, evidence: Any) -> float:
    def tokens(value: Any) -> set[str]:
        normalized = clean_text(value).casefold()
        values = {
            token for token in re.findall(r"[^\W\d_]{4,}", normalized, flags=re.UNICODE)
            if token not in STOPWORDS
        }
        return values

    target = tokens(context)
    observed = tokens(evidence)
    if not target or not observed:
        return 0.0
    return min(1.0, len(target & observed) / max(1, min(len(target), 60)))


def build_dossier_records(
    shortlist: dict,
    references: pd.DataFrame,
    config: dict,
) -> list[dict]:
    reference_map = references.set_index("reference_id").to_dict(orient="index")
    maximum_evidence = int(config["generation"]["maximum_evidence_items_per_reference"])
    maximum_description = int(config["generation"]["maximum_description_characters"])
    records = []
    for recommendation in shortlist["recommendations"].to_dict(orient="records"):
        reference_id = str(recommendation["reference_id"])
        if reference_id not in reference_map:
            raise TemplateGenerationError(f"Canonical reference missing: {reference_id}")
        metadata = reference_map[reference_id]
        start, end = project_year_bounds(metadata.get("project_year"))
        description = clean_text(metadata.get("service_nature"))[:maximum_description]
        if not description:
            raise TemplateGenerationError(f"Reference description missing: {reference_id}")
        evidence = shortlist["evidence"].loc[
            shortlist["evidence"]["reference_id"].astype(str).eq(reference_id)
        ].drop_duplicates(["citation_uri", "chunk_id"]).copy()
        grounded_title = mission_title(metadata)
        grounding_context = " ".join(
            clean_text(metadata.get(field))
            for field in ("service_nature", "offering", "client", "country")
        )
        evidence["grounding_overlap_score"] = evidence["evidence_excerpt"].map(
            lambda excerpt: max(
                lexical_grounding_score(grounding_context, excerpt),
                0.65 * lexical_grounding_score(grounded_title, excerpt)
                + 0.35 * lexical_grounding_score(grounding_context, excerpt),
            )
        )
        evidence = evidence.sort_values(
            ["grounding_overlap_score", "requirement_relevance_score", "citation_uri"],
            ascending=[False, False, True],
        )
        minimum_grounding = float(config["generation"]["minimum_evidence_grounding_overlap"])
        grounded_evidence = evidence.loc[evidence["grounding_overlap_score"].ge(minimum_grounding)]
        limited_evidence = grounded_evidence.empty
        evidence = (grounded_evidence if not grounded_evidence.empty else evidence.head(1)).head(maximum_evidence)
        coverage = shortlist["coverage"].loc[
            shortlist["coverage"]["reference_id"].astype(str).eq(reference_id)
        ].sort_values(["classification", "requirement_id"])
        records.append(
            {
                "reference_id": reference_id,
                "rank": int(recommendation["final_rank"]),
                "score": float(recommendation["base_score"]),
                "must_covered": int(recommendation["must_covered"]),
                "must_total": int(recommendation["must_total"]),
                "warnings": (
                    f"Couverture MUST {int(recommendation['must_covered'])}/{int(recommendation['must_total'])}."
                    if int(recommendation["must_covered"]) < int(recommendation["must_total"])
                    else "Pertinence textuelle des justificatifs à confirmer par un réviseur."
                    if limited_evidence
                    else ""
                ),
                "mission_title": grounded_title,
                "client": clean_text(metadata.get("client")) or "À confirmer",
                "country": clean_text(metadata.get("country")) or "À confirmer",
                "project_year": clean_text(metadata.get("project_year")) or "À confirmer",
                "start_year": start,
                "end_year": end,
                "offering": clean_text(metadata.get("offering")) or "À confirmer",
                "description": description,
                "recommendation_basis": (
                    "Sélection technique déterministe : couverture de "
                    f"{int(recommendation['must_covered'])}/{int(recommendation['must_total'])} exigences MUST."
                ),
                "evidence_review_required": limited_evidence,
                "coverage": coverage.to_dict(orient="records"),
                "evidence": evidence.to_dict(orient="records"),
            }
        )
    return records


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _format_cell(cell, *, size: float = 8.5, bold: bool = False, color: str = "000000", align=None) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Arial"
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(color)


def _set_repeat_header(row) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    value = OxmlElement("w:t")
    value.text = text
    run.extend([properties, value])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _set_table_widths(table, widths) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    total = sum(int(width.twips) for width in widths)
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(total))
    table_width.set(qn("w:type"), "dxa")
    table_indent = properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:w"), "0")
    table_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width.twips)))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
            tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_width.set(qn("w:w"), str(int(width.twips)))
            tc_width.set(qn("w:type"), "dxa")


def create_reference_docx(path: Path, records: list[dict], manifest: dict) -> None:
    from docx import Document
    from docx.enum.section import WD_ORIENT, WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt, RGBColor

    document = Document()
    landscape_section = document.sections[0]
    landscape_section.orientation = WD_ORIENT.LANDSCAPE
    landscape_section.page_width = Cm(29.7)
    landscape_section.page_height = Cm(21.0)
    landscape_section.top_margin = Cm(1.0)
    landscape_section.bottom_margin = Cm(1.0)
    landscape_section.left_margin = Cm(1.0)
    landscape_section.right_margin = Cm(1.0)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.5)

    draft = document.add_paragraph()
    draft.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = draft.add_run("DRAFT INTERNE — VALIDATION REQUISE")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(184, 22, 46)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(14)
    title_run = title.add_run("I.  Nos principales références")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)
    note = document.add_paragraph(
        "Sélection technique fondée sur les métadonnées canoniques et les justificatifs cités. "
        "Toute utilisation externe requiert une validation métier et sécurité."
    )
    note.paragraph_format.space_after = Pt(8)

    headers = ["#", "Intitulé du projet", "Client", "Pays", "Période", "Offre", "Score", "MUST", "Justificatifs"]
    summary = document.add_table(rows=1, cols=len(headers))
    summary.style = "Table Grid"
    widths = [Inches(x) for x in (0.35, 2.55, 1.45, 0.8, 0.8, 0.9, 0.55, 0.55, 1.65)]
    for index, header in enumerate(headers):
        summary.cell(0, index).text = header
        _set_cell_shading(summary.cell(0, index), "F9425B")
        _format_cell(summary.cell(0, index), size=8, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_repeat_header(summary.rows[0])
    for record in records:
        row = summary.add_row().cells
        values = [
            record["rank"], record["mission_title"], record["client"], record["country"],
            record["project_year"], record["offering"], f"{record['score']:.3f}",
            f"{record['must_covered']}/{record['must_total']}",
            f"{len(record['evidence'])} citation(s)" + (" — À revoir" if record["evidence_review_required"] else ""),
        ]
        for index, value in enumerate(values):
            row[index].text = str(value)
            _format_cell(
                row[index], size=8.1, bold=index == 0,
                align=WD_ALIGN_PARAGRAPH.CENTER if index in {0, 3, 4, 6, 7, 8} else WD_ALIGN_PARAGRAPH.LEFT,
            )
    _set_table_widths(summary, widths)

    portrait = document.add_section(WD_SECTION.NEW_PAGE)
    portrait.orientation = WD_ORIENT.PORTRAIT
    portrait.page_width = Cm(21.0)
    portrait.page_height = Cm(29.7)
    portrait.top_margin = Cm(1.5)
    portrait.bottom_margin = Cm(1.5)
    portrait.left_margin = Cm(1.8)
    portrait.right_margin = Cm(1.8)
    heading = document.add_paragraph()
    heading_run = heading.add_run("II.  Annexes")
    heading_run.bold = True
    heading_run.font.name = "Arial"
    heading_run.font.size = Pt(22)
    subheading = document.add_paragraph()
    sub_run = subheading.add_run("1.  Description détaillée des références et justificatifs")
    sub_run.bold = True
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(15)

    for position, record in enumerate(records, start=1):
        # Use an explicit break paragraph between references. LibreOffice can
        # mis-layout a table that immediately follows a pageBreakBefore
        # paragraph (observed as a clipped first column on some records).
        if position > 1:
            document.add_page_break()
        ref_heading = document.add_paragraph()
        ref_heading.paragraph_format.space_before = Pt(12)
        ref_heading.paragraph_format.space_after = Pt(6)
        ref_run = ref_heading.add_run(f"Référence N°{position}")
        ref_run.bold = True
        ref_run.font.name = "Arial"
        ref_run.font.size = Pt(14)
        fields = [
            ("Nom de la mission", record["mission_title"]),
            ("Pays", record["country"]),
            ("Nom de l’Autorité Contractante", record["client"]),
            ("Date de démarrage", record["start_year"]),
            ("Date d’achèvement", record["end_year"]),
            ("Offre", record["offering"]),
            ("Description du projet", record["description"]),
            ("Motif de sélection", record["recommendation_basis"] or "Classement technique déterministe."),
            ("Avertissement", record["warnings"] or "Aucun avertissement technique."),
        ]
        detail = document.add_table(rows=len(fields), cols=2)
        detail.style = "Table Grid"
        for row_index, (label, value) in enumerate(fields):
            detail.cell(row_index, 0).text = label
            detail.cell(row_index, 1).text = value
            _set_cell_shading(detail.cell(row_index, 0), "F5A0AE")
            _format_cell(detail.cell(row_index, 0), size=9, bold=False)
            _format_cell(detail.cell(row_index, 1), size=9)
        _set_table_widths(detail, [Inches(1.72), Inches(5.0)])

        coverage_heading = document.add_paragraph()
        coverage_heading.paragraph_format.space_before = Pt(8)
        coverage_heading.paragraph_format.space_after = Pt(4)
        coverage_heading.add_run("Couverture des exigences").bold = True
        coverage_table = document.add_table(rows=1, cols=3)
        coverage_table.style = "Table Grid"
        for index, value in enumerate(("Classe", "Couvert", "Exigence")):
            coverage_table.cell(0, index).text = value
            _set_cell_shading(coverage_table.cell(0, index), "F9425B")
            _format_cell(coverage_table.cell(0, index), size=8, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_repeat_header(coverage_table.rows[0])
        for coverage in record["coverage"]:
            cells = coverage_table.add_row().cells
            cells[0].text = clean_text(coverage.get("classification"))
            cells[1].text = "Oui" if bool(coverage.get("covered")) else "Non"
            cells[2].text = clean_text(coverage.get("requirement_text"))
            for index, cell in enumerate(cells):
                _format_cell(cell, size=8.2, align=WD_ALIGN_PARAGRAPH.CENTER if index < 2 else WD_ALIGN_PARAGRAPH.LEFT)
        _set_table_widths(coverage_table, [Inches(0.9), Inches(0.7), Inches(5.12)])

        evidence_heading = document.add_paragraph()
        evidence_heading.paragraph_format.space_before = Pt(8)
        evidence_heading.paragraph_format.space_after = Pt(3)
        evidence_heading.add_run("Justificatifs cités").bold = True
        for evidence in record["evidence"]:
            paragraph = document.add_paragraph(style=None)
            paragraph.paragraph_format.left_indent = Cm(0.4)
            paragraph.paragraph_format.first_line_indent = Cm(-0.4)
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(f"• {clean_text(evidence.get('citation_label'))} — ")
            run.bold = True
            excerpt = clean_text(evidence.get("evidence_excerpt"))[:260]
            paragraph.add_run(excerpt + " ")
            _add_hyperlink(paragraph, "Ouvrir la source", clean_text(evidence.get("citation_uri")))

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "DRAFT INTERNE — VALIDATION MÉTIER ET SÉCURITÉ REQUISE"
        for run in footer.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(110, 110, 110)
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def create_reference_pdf(path: Path, records: list[dict], manifest: dict) -> None:
    import reportlab
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        BaseDocTemplate, Frame, NextPageTemplate, PageBreak, PageTemplate,
        Paragraph, Spacer, Table, TableStyle,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Embed a Unicode-capable font shipped with ReportLab so the PDF renders
    # consistently on Colab, Windows, and reviewer machines.
    font_root = Path(reportlab.__file__).resolve().parent / "fonts"
    pdfmetrics.registerFont(TTFont("DevoteamSans", str(font_root / "Vera.ttf")))
    pdfmetrics.registerFont(TTFont("DevoteamSans-Bold", str(font_root / "VeraBd.ttf")))
    pdfmetrics.registerFontFamily(
        "DevoteamSans",
        normal="DevoteamSans",
        bold="DevoteamSans-Bold",
        italic="DevoteamSans",
        boldItalic="DevoteamSans-Bold",
    )

    styles = getSampleStyleSheet()
    body = ParagraphStyle("BodyFR", parent=styles["BodyText"], fontName="DevoteamSans", fontSize=8.3, leading=10.2)
    small = ParagraphStyle("SmallFR", parent=body, fontSize=7.2, leading=8.5)
    title = ParagraphStyle("TitleFR", parent=styles["Heading1"], fontName="DevoteamSans-Bold", fontSize=19, leading=22)
    heading = ParagraphStyle("HeadingFR", parent=styles["Heading2"], fontName="DevoteamSans-Bold", fontSize=13, leading=16)
    draft = ParagraphStyle("Draft", parent=small, alignment=TA_RIGHT, textColor=colors.HexColor("#B8162E"))
    landscape_frame = Frame(10 * mm, 10 * mm, landscape(A4)[0] - 20 * mm, landscape(A4)[1] - 20 * mm, id="landscape")
    portrait_frame = Frame(16 * mm, 16 * mm, A4[0] - 32 * mm, A4[1] - 32 * mm, id="portrait")
    document = BaseDocTemplate(str(path), pageTemplates=[
        PageTemplate(id="summary", pagesize=landscape(A4), frames=[landscape_frame]),
        PageTemplate(id="details", pagesize=A4, frames=[portrait_frame]),
    ])
    story = [
        Paragraph("DRAFT INTERNE — VALIDATION REQUISE", draft),
        Paragraph("I. Nos principales références", title),
        Spacer(1, 5 * mm),
    ]
    summary_data = [[Paragraph(escape(value), small) for value in ("#", "Intitulé", "Client", "Pays", "Période", "Offre", "Score", "MUST", "Justificatifs")]]
    for record in records:
        summary_data.append([
            str(record["rank"]), Paragraph(escape(record["mission_title"]), small),
            Paragraph(escape(record["client"]), small), Paragraph(escape(record["country"]), small),
            record["project_year"], Paragraph(escape(record["offering"]), small),
            f"{record['score']:.3f}", f"{record['must_covered']}/{record['must_total']}",
            f"{len(record['evidence'])} citation(s)" + (" - À revoir" if record["evidence_review_required"] else ""),
        ])
    summary_table = Table(
        summary_data,
        colWidths=[value * mm for value in (8, 60, 38, 25, 25, 29, 19, 18, 30)],
        repeatRows=1,
    )
    summary_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F9425B")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), "DevoteamSans"),
        ("FONTNAME", (0, 0), (-1, 0), "DevoteamSans-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("ALIGN", (3, 1), (-1, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.extend([summary_table, NextPageTemplate("details"), PageBreak(), Paragraph("II. Annexes", title), Paragraph("1. Description détaillée des références et justificatifs", heading), Spacer(1, 4 * mm)])
    for position, record in enumerate(records, start=1):
        story.append(Paragraph(f"Référence N°{position}", heading))
        details = [
            ("Nom de la mission", record["mission_title"]), ("Pays", record["country"]),
            ("Nom de l’Autorité Contractante", record["client"]), ("Date de démarrage", record["start_year"]),
            ("Date d’achèvement", record["end_year"]), ("Offre", record["offering"]),
            ("Description du projet", record["description"]), ("Motif de sélection", record["recommendation_basis"]),
            ("Avertissement", record["warnings"] or "Aucun avertissement technique."),
        ]
        detail_data = [[Paragraph(escape(label), body), Paragraph(escape(value), body)] for label, value in details]
        table = Table(detail_data, colWidths=[43 * mm, 126 * mm])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F5A0AE")),
            ("FONTNAME", (0, 0), (-1, -1), "DevoteamSans"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.extend([table, Spacer(1, 4 * mm), Paragraph("Couverture des exigences", heading)])
        coverage_data = [[Paragraph("Classe", small), Paragraph("Couvert", small), Paragraph("Exigence", small)]]
        for coverage in record["coverage"]:
            coverage_data.append([
                Paragraph(escape(clean_text(coverage.get("classification"))), small),
                "Oui" if bool(coverage.get("covered")) else "Non",
                Paragraph(escape(clean_text(coverage.get("requirement_text"))), small),
            ])
        coverage_table = Table(coverage_data, colWidths=[24 * mm, 18 * mm, 127 * mm], repeatRows=1)
        coverage_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F9425B")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, -1), "DevoteamSans"),
            ("FONTNAME", (0, 0), (-1, 0), "DevoteamSans-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ALIGN", (0, 0), (1, -1), "CENTER"),
            ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.extend([coverage_table, Spacer(1, 4 * mm), Paragraph("Justificatifs cités", heading)])
        for evidence in record["evidence"]:
            label = escape(clean_text(evidence.get("citation_label")))
            excerpt = escape(clean_text(evidence.get("evidence_excerpt"))[:260])
            url = escape(clean_text(evidence.get("citation_uri")), {'"': '&quot;'})
            story.append(Paragraph(f"• <b>{label}</b> — {excerpt} <link href=\"{url}\">Ouvrir la source</link>", body))
            story.append(Spacer(1, 1.5 * mm))
        if position < len(records):
            story.append(PageBreak())
    path.parent.mkdir(parents=True, exist_ok=True)
    document.build(story)


def _build_report(manifest: dict, records: list[dict]) -> str:
    lines = [
        "# Phase 8 — Controlled reference dossier report", "",
        f"**Status:** {manifest['status']}  ",
        f"**Opportunity:** `{manifest['opportunity_id']}`  ",
        f"**Selection mode:** `{manifest['selection_mode']}`  ", "",
        "## Selected references", "",
    ]
    for record in records:
        lines.append(
            f"{record['rank']}. **{record['client']}** — {record['mission_title']} "
            f"(score {record['score']:.3f}; MUST {record['must_covered']}/{record['must_total']})."
        )
        if record["warnings"]:
            lines.append(f"   - Avertissement : {record['warnings']}")
    lines.extend([
        "", "## Controls", "",
        "- The supplied template fingerprint and structure were verified.",
        "- All generated factual content comes from canonical reference metadata.",
        "- Every selected reference has at least one exact Drive/page citation.",
        "- Source documents and page images were not copied into the output.",
        "- No LLM, translation API, or automatic client-delivery action was used.",
        "- Human business and security validation remains mandatory.", "",
    ])
    return "\n".join(lines)


def run_phase8(
    *,
    phase4_root: Path,
    phase7_root: Path,
    template_path: Path,
    phase7_config: dict,
    phase8_config: dict,
    output_root: Path,
) -> tuple[Path, dict]:
    template = verify_reference_template(template_path, phase8_config)
    shortlist = load_selected_shortlist(phase7_root, phase7_config, phase8_config)
    references_path = phase4_root / "reference_catalog.parquet"
    if sha256_file(references_path) != phase8_config["input"]["expected_references_sha256"]:
        raise AssertionError("Phase 4 reference catalogue changed")
    references = pd.read_parquet(references_path)
    records = build_dossier_records(shortlist, references, phase8_config)
    opportunity_id = shortlist["phase7_manifest"]["opportunity_id"]
    run_root = output_root / opportunity_id / phase8_config["pipeline_version"]
    success_path = run_root / phase8_config["output"]["success_marker"]
    if success_path.exists():
        verification = verify_phase8_run(run_root, phase8_config)
        return run_root, verification["manifest"]
    run_root.mkdir(parents=True, exist_ok=True)
    evidence_review_count = sum(int(record["evidence_review_required"]) for record in records)
    if shortlist["phase7_manifest"]["status"] == "TECHNICAL_PASS_SAMPLE_ONLY":
        status = (
            "TECHNICAL_PASS_SAMPLE_ONLY_WITH_EVIDENCE_WARNINGS"
            if evidence_review_count
            else "TECHNICAL_PASS_SAMPLE_ONLY"
        )
    else:
        status = "DRAFT_BLOCKED_EVIDENCE_REVIEW" if evidence_review_count else "DRAFT_READY_FOR_HUMAN_VALIDATION"
    manifest = {
        "schema_version": 1,
        "phase": 8,
        "pipeline_version": phase8_config["pipeline_version"],
        "completed_at_utc": utc_now(),
        "status": status,
        "opportunity_id": opportunity_id,
        "phase7_manifest_sha256": shortlist["phase7_manifest_sha256"],
        "template_sha256": template["sha256"],
        "template_structure_verified": True,
        "selection_mode": shortlist["selection_mode"],
        "selected_references": len(records),
        "selected_reference_ids": [record["reference_id"] for record in records],
        "citation_rows": sum(len(record["evidence"]) for record in records),
        "citation_coverage": 1.0,
        "references_requiring_evidence_review": evidence_review_count,
        "generation_mode": phase8_config["generation"]["mode"],
        "external_llm_calls": 0,
        "external_translation_calls": 0,
        "raw_opportunity_log_rows": 0,
        "source_documents_copied": False,
        "source_page_images_embedded": False,
        "automatic_client_delivery": False,
        "production_promotion_allowed": False,
        "business_validation_required": True,
        "security_validation_required": True,
    }
    docx_path = run_root / phase8_config["output"]["docx_name"]
    pdf_path = run_root / phase8_config["output"]["pdf_name"]
    create_reference_docx(docx_path, records, manifest)
    create_reference_pdf(pdf_path, records, manifest)
    _atomic_json(run_root / "REFERENCE_DOSSIER_DATA.json", records)
    _atomic_json(run_root / "PHASE_8_MANIFEST.json", manifest)
    _atomic_text(run_root / "PHASE_8_REPORT.md", _build_report(manifest, records))
    hashed = [
        path for path in run_root.iterdir()
        if path.is_file() and path.name not in {"SHA256SUMS.txt", phase8_config["output"]["success_marker"]}
    ]
    sums = "".join(f"{sha256_file(path)}  {path.name}\n" for path in sorted(hashed, key=lambda item: item.name))
    sums_path = run_root / "SHA256SUMS.txt"
    _atomic_text(sums_path, sums)
    manifest_path = run_root / "PHASE_8_MANIFEST.json"
    _atomic_json(
        success_path,
        {
            "status": "COMPLETE_REPRODUCIBLE_REFERENCE_DOSSIER",
            "created_at_utc": utc_now(),
            "pipeline_version": phase8_config["pipeline_version"],
            "opportunity_id": opportunity_id,
            "manifest_sha256": sha256_file(manifest_path),
            "sha256sums_sha256": sha256_file(sums_path),
        },
    )
    verify_phase8_run(run_root, phase8_config)
    return run_root, manifest


def verify_phase8_run(run_root: Path, config: dict) -> dict:
    success_path = run_root / config["output"]["success_marker"]
    manifest_path = run_root / "PHASE_8_MANIFEST.json"
    sums_path = run_root / "SHA256SUMS.txt"
    for path in (success_path, manifest_path, sums_path):
        if not path.exists():
            raise FileNotFoundError(path)
    success = json.loads(success_path.read_text(encoding="utf-8"))
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    if success["manifest_sha256"] != sha256_file(manifest_path):
        raise AssertionError("Phase 8 manifest hash mismatch")
    if success["sha256sums_sha256"] != sha256_file(sums_path):
        raise AssertionError("Phase 8 checksum-file hash mismatch")
    for line in sums_path.read_text(encoding="utf-8").splitlines():
        expected, name = line.split("  ", 1)
        if sha256_file(run_root / name) != expected:
            raise AssertionError(f"Phase 8 output changed: {name}")
    for name in (config["output"]["docx_name"], config["output"]["pdf_name"], "REFERENCE_DOSSIER_DATA.json"):
        if not (run_root / name).exists():
            raise FileNotFoundError(run_root / name)
    prohibited = ("external_llm_calls", "external_translation_calls", "raw_opportunity_log_rows")
    if any(int(manifest.get(field, -1)) != 0 for field in prohibited):
        raise AssertionError("Phase 8 provider/security boundary failed")
    if manifest["source_documents_copied"] or manifest["source_page_images_embedded"]:
        raise AssertionError("Phase 8 source-copy boundary failed")
    if manifest["automatic_client_delivery"] or manifest["production_promotion_allowed"]:
        raise AssertionError("Phase 8 delivery/promotion boundary failed")
    if manifest["citation_coverage"] != 1.0 or not manifest["template_structure_verified"]:
        raise AssertionError("Phase 8 evidence/template gate failed")
    if int(manifest["selected_references"]) < 1:
        raise AssertionError("Phase 8 generated an empty dossier")
    return {"success": success, "manifest": manifest}
