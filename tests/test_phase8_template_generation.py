from __future__ import annotations

import json
import os
from pathlib import Path

import openpyxl
import pandas as pd
import pytest
import yaml
from docx import Document

import devoteam_reference_ai.phase8_template_generation as phase8
from devoteam_reference_ai.phase8_template_generation import (
    TemplateGenerationError,
    build_dossier_records,
    clean_text,
    create_reference_docx,
    create_reference_pdf,
    load_phase8_config,
    load_selected_shortlist,
    lexical_grounding_score,
    mission_title,
    project_year_bounds,
    verify_reference_template,
)


ROOT = Path(__file__).resolve().parents[1]
CONFIG_PATH = ROOT / "config" / "phase8_template_generation.yaml"
REAL_TEMPLATE = Path(
    os.environ.get(
        "DEVOTEAM_PHASE8_TEMPLATE",
        ROOT / "human_inputs" / "phase8" / "REFERENCE_TEMPLATE.docx",
    )
).resolve()


@pytest.fixture()
def config() -> dict:
    return load_phase8_config(CONFIG_PATH)


def test_configuration_loads_deterministic_secure_mode(config):
    assert config["generation"]["mode"] == "deterministic_template_derived"
    assert config["generation"]["external_llm_enabled"] is False
    assert config["generation"]["automatic_client_delivery"] is False


@pytest.mark.parametrize(
    "section,field",
    [
        ("generation", "external_llm_enabled"),
        ("generation", "external_translation_enabled"),
        ("generation", "automatic_client_delivery"),
        ("generation", "copy_source_documents_into_output"),
        ("generation", "embed_source_page_images"),
        ("security", "raw_opportunity_logging_allowed"),
        ("security", "source_mutation_allowed"),
        ("security", "hidden_content_generation_allowed"),
    ],
)
def test_security_defaults_cannot_be_weakened(tmp_path, config, section, field):
    changed = json.loads(json.dumps(config))
    changed[section][field] = True
    path = tmp_path / "bad.yaml"
    path.write_text(yaml.safe_dump(changed), encoding="utf-8")
    with pytest.raises(ValueError):
        load_phase8_config(path)


def test_template_fingerprint_and_structure_are_exact(config):
    result = verify_reference_template(REAL_TEMPLATE, config)
    assert result["sha256"] == config["input"]["expected_template_sha256"]
    assert result["reference_slots"] == 17


def test_template_tampering_is_detected(tmp_path, config):
    changed = tmp_path / "changed.docx"
    changed.write_bytes(REAL_TEMPLATE.read_bytes() + b"tamper")
    with pytest.raises(AssertionError):
        verify_reference_template(changed, config)


@pytest.mark.parametrize(
    "value,expected",
    [
        ("2021", ("2021", "2021")),
        ("2019-2022", ("2019", "2022")),
        ("2024-Présent", ("2024", "2024")),
        ("", ("À confirmer", "À confirmer")),
    ],
)
def test_project_year_bounds(value, expected):
    assert project_year_bounds(value) == expected


def test_mission_title_uses_grounded_service_text():
    row = {"service_nature": "Élaboration du schéma directeur du système d’information. Deuxième phrase.", "offering": "SDSI", "client": "Banque"}
    assert mission_title(row) == "Élaboration du schéma directeur du système d’information."


def test_mission_title_stops_before_phase_or_bullet_details():
    row = {
        "service_nature": "Accompagnement de la banque dans sa transformation digitale PHASE 1 : Diagnostic • Livrable",
        "offering": "AMOA",
        "client": "Banque",
    }
    assert mission_title(row) == "Accompagnement de la banque dans sa transformation digitale"


def test_mission_title_has_deterministic_fallback():
    assert mission_title({"service_nature": "Court", "offering": "Audit", "client": "Banque A"}) == "Audit — Banque A"


def test_lexical_grounding_prefers_project_evidence_over_contract_boilerplate():
    context = "Accompagnement au pilotage du portefeuille de projets SI et installation du PMO"
    relevant = "Le présent contrat a pour objet l'accompagnement au pilotage du portefeuille projets SI et le PMO."
    irrelevant = "Les pénalités ne pourront excéder dix pour cent du montant total du contrat."
    assert lexical_grounding_score(context, relevant) > lexical_grounding_score(context, irrelevant)


def test_evidence_warning_is_explicit_when_no_citation_meets_threshold(config):
    changed = json.loads(json.dumps(config))
    changed["generation"]["minimum_evidence_grounding_overlap"] = 1.0
    records = build_dossier_records(_shortlist(), _references(), changed)
    assert records[0]["evidence_review_required"] is True
    assert "Pertinence textuelle" in records[0]["warnings"]


def _recommendations() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "final_rank": 1,
                "reference_id": "REF1",
                "base_score": 0.9,
                "must_covered": 2,
                "must_total": 2,
                "warnings": "",
                "recommendation_basis": "Classement technique déterministe.",
            },
            {
                "final_rank": 2,
                "reference_id": "REF2",
                "base_score": 0.7,
                "must_covered": 1,
                "must_total": 2,
                "warnings": "MUST coverage 1/2",
                "recommendation_basis": "Classement technique déterministe.",
            },
        ]
    )


def _evidence() -> pd.DataFrame:
    rows = []
    for reference_id in ("REF1", "REF2"):
        rows.append(
            {
                "reference_id": reference_id,
                "requirement_id": "REQ1",
                "requirement_relevance_score": 0.8,
                "citation_uri": f"https://drive.google.com/{reference_id}#page=1",
                "citation_label": f"{reference_id}.pdf p.1",
                "evidence_excerpt": "Extrait de preuve vérifiable.",
                "source_sha256": "a" * 64,
                "chunk_text_sha256": "b" * 64,
                "chunk_id": f"CHUNK-{reference_id}",
            }
        )
    return pd.DataFrame(rows)


def _coverage() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {"reference_id": reference_id, "requirement_id": "REQ1", "classification": "MUST", "requirement_text": "Une exigence obligatoire", "covered": True}
            for reference_id in ("REF1", "REF2")
        ]
    )


def _write_phase7_files(root: Path) -> None:
    root.mkdir(parents=True)
    (root / "PHASE_7_MANIFEST.json").write_text(
        json.dumps({"status": "TEST", "opportunity_id": "OPP"}), encoding="utf-8"
    )
    _recommendations().to_parquet(root / "recommendations.parquet", index=False)
    _evidence().to_parquet(root / "evidence_matrix.parquet", index=False)
    _coverage().to_csv(root / "requirement_coverage.csv", index=False, encoding="utf-8-sig")


def test_sample_shortlist_uses_bounded_top_n(tmp_path, config, monkeypatch):
    phase7_root = tmp_path / "phase7"
    _write_phase7_files(phase7_root)
    monkeypatch.setattr(
        phase8,
        "verify_phase7_run",
        lambda *_: {"manifest": {"status": "TECHNICAL_PASS_SAMPLE_ONLY", "opportunity_id": "OPP"}},
    )
    changed = json.loads(json.dumps(config))
    changed["review"]["synthetic_sample_top_n"] = 1
    result = load_selected_shortlist(phase7_root, {}, changed)
    assert result["recommendations"]["reference_id"].tolist() == ["REF1"]
    assert result["selection_mode"] == "SYNTHETIC_DEVELOPMENT_FIXTURE_TOP_N"


def test_real_shortlist_requires_complete_review(tmp_path, config, monkeypatch):
    phase7_root = tmp_path / "phase7"
    _write_phase7_files(phase7_root)
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Ranked References"
    sheet.append(["reference_id", "reviewer_decision"])
    sheet.append(["REF1", "SHORTLIST"])
    sheet.append(["REF2", "PENDING"])
    workbook.save(phase7_root / "RECOMMENDATION_REVIEW.xlsx")
    monkeypatch.setattr(
        phase8,
        "verify_phase7_run",
        lambda *_: {"manifest": {"status": "READY_FOR_BUSINESS_SHORTLIST_REVIEW", "opportunity_id": "OPP"}},
    )
    with pytest.raises(TemplateGenerationError, match="Every real recommendation"):
        load_selected_shortlist(phase7_root, {}, config)


def test_real_shortlist_uses_only_human_approved_rows(tmp_path, config, monkeypatch):
    phase7_root = tmp_path / "phase7"
    _write_phase7_files(phase7_root)
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Ranked References"
    sheet.append(["reference_id", "reviewer_decision"])
    sheet.append(["REF1", "SHORTLIST"])
    sheet.append(["REF2", "REJECT"])
    workbook.save(phase7_root / "RECOMMENDATION_REVIEW.xlsx")
    monkeypatch.setattr(
        phase8,
        "verify_phase7_run",
        lambda *_: {"manifest": {"status": "READY_FOR_BUSINESS_SHORTLIST_REVIEW", "opportunity_id": "OPP"}},
    )
    result = load_selected_shortlist(phase7_root, {}, config)
    assert result["recommendations"]["reference_id"].tolist() == ["REF1"]
    assert set(result["evidence"]["reference_id"]) == {"REF1"}
    assert result["selection_mode"] == "HUMAN_APPROVED_SHORTLIST"


def test_missing_selected_evidence_fails_closed(tmp_path, config, monkeypatch):
    phase7_root = tmp_path / "phase7"
    _write_phase7_files(phase7_root)
    _evidence().loc[lambda frame: frame["reference_id"].eq("REF1")].to_parquet(
        phase7_root / "evidence_matrix.parquet", index=False
    )
    monkeypatch.setattr(
        phase8,
        "verify_phase7_run",
        lambda *_: {"manifest": {"status": "TECHNICAL_PASS_SAMPLE_ONLY", "opportunity_id": "OPP"}},
    )
    with pytest.raises(TemplateGenerationError, match="Every selected reference"):
        load_selected_shortlist(phase7_root, {}, config)


def _shortlist() -> dict:
    return {
        "recommendations": _recommendations().head(1),
        "evidence": _evidence().loc[lambda frame: frame["reference_id"].eq("REF1")],
        "coverage": _coverage().loc[lambda frame: frame["reference_id"].eq("REF1")],
    }


def _references() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "reference_id": "REF1",
                "client": "Banque A",
                "country": "Tunisie",
                "project_year": "2021-2022",
                "offering": "SDSI",
                "service_nature": "Élaboration du schéma directeur du système d’information. Définition de la feuille de route.",
            }
        ]
    )


def test_dossier_records_are_grounded_and_cited(config):
    records = build_dossier_records(_shortlist(), _references(), config)
    assert records[0]["client"] == "Banque A"
    assert records[0]["start_year"] == "2021"
    assert records[0]["end_year"] == "2022"
    assert records[0]["evidence"][0]["citation_uri"].startswith("https://drive.google.com/")
    assert "Sélection technique déterministe" in records[0]["recommendation_basis"]


def test_docx_preserves_template_derived_structure_and_draft_gate(tmp_path, config):
    records = build_dossier_records(_shortlist(), _references(), config)
    path = tmp_path / "dossier.docx"
    create_reference_docx(path, records, {"status": "TEST"})
    document = Document(path)
    assert len(document.tables) == 3
    assert "Nos principales références" in " ".join(paragraph.text for paragraph in document.paragraphs)
    assert "DRAFT INTERNE" in " ".join(paragraph.text for paragraph in document.paragraphs)
    assert "hyperlink" in path.read_bytes().decode("latin-1", errors="ignore") or path.stat().st_size > 5000


def test_pdf_is_created_with_real_pdf_signature(tmp_path, config):
    records = build_dossier_records(_shortlist(), _references(), config)
    path = tmp_path / "dossier.pdf"
    create_reference_pdf(path, records, {"status": "TEST"})
    assert path.read_bytes().startswith(b"%PDF")
    assert path.stat().st_size > 3000


def test_clean_text_normalizes_whitespace_and_nan():
    assert clean_text("  a\n b  ") == "a b"
    assert clean_text(float("nan")) == ""
